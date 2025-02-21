from PyQt6.QtWidgets import *
from PyQt6.QtCore import *
from PyQt6.QtGui import QAction, QPainter, QFont
from PyQt6.QtPrintSupport import QPrintPreviewDialog, QPrinter
import sys
import json
from pathlib import Path
from ai_analyzer import MedicalAnalyzer
from datetime import datetime
import re
import platform

class MedicalAssistant(QMainWindow):
    def __init__(self):
        super().__init__()
        self.setWindowTitle("AI医疗助手 - AI安全工坊出品（微信公众号搜索关注）")
        self.setMinimumSize(1200, 800)
        
        # 初始化AI分析器
        self.analyzer = MedicalAnalyzer()
        
        # 初始化数据存储 - 移到这里，在创建界面之前
        self.init_storage()
        
        # 创建主窗口部件
        main_widget = QWidget()
        self.setCentralWidget(main_widget)
        
        # 创建水平布局
        main_layout = QHBoxLayout(main_widget)
        
        # 左侧面板
        left_panel = QVBoxLayout()
        left_panel.addWidget(self.create_model_selector())
        left_panel.addWidget(self.create_patient_info())
        left_panel.addWidget(self.create_symptom_tags())  # 新增症状标签系统
        left_panel.addWidget(self.create_smart_inquiry())  # 新增智能问诊
        
        # 中间面板
        middle_panel = QVBoxLayout()
        middle_panel.addWidget(self.create_input_area())
        middle_panel.addWidget(self.create_output_area())
        middle_panel.addLayout(self.create_buttons())
        
        # 右侧面板
        right_panel = QVBoxLayout()
        right_panel.addWidget(self.create_medication_reminder())  # 新增用药提醒
        right_panel.addWidget(self.create_health_trends())  # 新增健康趋势图
        right_panel.addWidget(self.create_prescription_manager())  # 新增处方管理功能
        
        # 将三个面板添加到主布局
        main_layout.addLayout(left_panel, 2)
        main_layout.addLayout(middle_panel, 4)
        main_layout.addLayout(right_panel, 2)
        
        # 创建菜单栏和状态栏
        self.create_menu()
        self.statusBar = self.statusBar()
        self.progress_bar = QProgressBar()
        self.progress_bar.setVisible(False)
        self.statusBar.addPermanentWidget(self.progress_bar)

        # 显示引导教程
        self.show_tutorial()  # 确保在所有组件初始化后调用

    def init_storage(self):
        """初始化数据存储"""
        import sqlite3
        
        try:
            # 创建数据库连接
            self.db = sqlite3.connect('medical.db')
            self.cursor = self.db.cursor()
            
            # 创建必要的表
            self.cursor.execute('''
                CREATE TABLE IF NOT EXISTS medical_records (
                    id INTEGER PRIMARY KEY AUTOINCREMENT,
                    timestamp TEXT,
                    patient_info TEXT,
                    symptoms TEXT,
                    diagnosis TEXT
                )
            ''')
            
            self.cursor.execute('''
                CREATE TABLE IF NOT EXISTS health_trends (
                    id INTEGER PRIMARY KEY AUTOINCREMENT,
                    timestamp TEXT,
                    type TEXT,
                    value REAL
                )
            ''')
            
            self.cursor.execute('''
                CREATE TABLE IF NOT EXISTS medication_reminders (
                    id INTEGER PRIMARY KEY AUTOINCREMENT,
                    medicine_name TEXT,
                    dosage TEXT,
                    time TEXT,
                    notes TEXT
                )
            ''')
            
            # 创建处方相关表
            self.cursor.execute('''
                CREATE TABLE IF NOT EXISTS prescriptions (
                    id INTEGER PRIMARY KEY AUTOINCREMENT,
                    prescription_no TEXT UNIQUE,  -- 处方编号
                    type TEXT,                    -- 处方类型（普通、急诊、儿科等）
                    category TEXT,                -- 处方分类（西药、中药、中成药）
                    date TEXT,                    -- 开具日期
                    validity TEXT,                -- 有效期
                    patient_name TEXT,            -- 患者姓名
                    patient_gender TEXT,          -- 患者性别
                    patient_age INTEGER,          -- 患者年龄
                    patient_weight REAL,          -- 患者体重
                    medical_insurance TEXT,       -- 医保类型
                    diagnosis TEXT,               -- 诊断结果
                    doctor_name TEXT,             -- 医师姓名
                    doctor_title TEXT,            -- 医师职称
                    hospital_name TEXT,           -- 医疗机构名称
                    department TEXT,              -- 科室
                    status TEXT,                  -- 状态（未调配、已调配、已发药等）
                    notes TEXT                    -- 备注
                )
            ''')
            
            self.cursor.execute('''
                CREATE TABLE IF NOT EXISTS prescription_items (
                    id INTEGER PRIMARY KEY AUTOINCREMENT,
                    prescription_id INTEGER,       -- 关联处方ID
                    medicine_name TEXT,            -- 药品名称
                    specification TEXT,            -- 规格
                    dosage TEXT,                   -- 用法用量
                    frequency TEXT,                -- 频次
                    quantity REAL,                 -- 数量
                    unit TEXT,                     -- 单位
                    usage_method TEXT,             -- 用药方法
                    notes TEXT,                    -- 用药说明
                    FOREIGN KEY (prescription_id) REFERENCES prescriptions (id)
                )
            ''')
            
            self.db.commit()
            
        except Exception as e:
            QMessageBox.critical(self, "错误", f"初始化数据库失败: {str(e)}")
            raise

    def create_model_selector(self):
        group = QGroupBox("选择AI模型")
        group_layout = QHBoxLayout()
        
        self.model_combo = QComboBox()
        self.model_combo.addItems(["OpenAI", "DeepSeek", "双模型分析"])
        group_layout.addWidget(self.model_combo)
        
        group.setLayout(group_layout)
        return group

    def create_patient_info(self):
        group = QGroupBox("患者信息")
        group_layout = QFormLayout()
        
        # 年龄输入
        self.age_input = QSpinBox()
        self.age_input.setRange(0, 120)
        group_layout.addRow("年龄:", self.age_input)
        
        # 性别选择
        self.gender_combo = QComboBox()
        self.gender_combo.addItems(["男", "女"])
        group_layout.addRow("性别:", self.gender_combo)
        
        # 其他基本信息
        self.height_input = QSpinBox()
        self.height_input.setRange(0, 250)
        self.height_input.setSuffix(" cm")
        group_layout.addRow("身高:", self.height_input)
        
        self.weight_input = QDoubleSpinBox()
        self.weight_input.setRange(0, 200)
        self.weight_input.setSuffix(" kg")
        group_layout.addRow("体重:", self.weight_input)
        
        group.setLayout(group_layout)
        return group

    def create_input_area(self):
        group = QGroupBox("症状描述")
        group_layout = QVBoxLayout()
        
        # 症状输入框
        self.symptoms_text = QTextEdit()
        self.symptoms_text.setPlaceholderText("请详细描述您的症状...")
        group_layout.addWidget(self.symptoms_text)
        
        # 常见症状快速选择
        common_symptoms = QHBoxLayout()
        symptoms = ["发热", "头痛", "咳嗽", "腹痛", "恶心", "乏力"]
        for symptom in symptoms:
            btn = QPushButton(symptom)
            btn.clicked.connect(lambda checked, s=symptom: self.add_symptom(s))
            common_symptoms.addWidget(btn)
        
        group_layout.addLayout(common_symptoms)
        group.setLayout(group_layout)
        return group

    def create_output_area(self):
        group = QGroupBox("诊断结果")
        group_layout = QVBoxLayout()
        
        self.output_text = QTextEdit()
        self.output_text.setReadOnly(True)
        group_layout.addWidget(self.output_text)
        
        group.setLayout(group_layout)
        return group

    def create_buttons(self):
        button_layout = QHBoxLayout()
        
        # 分析按钮
        self.analyze_btn = QPushButton("开始分析")
        self.analyze_btn.clicked.connect(self.analyze_symptoms)
        button_layout.addWidget(self.analyze_btn)
        
        # 清空按钮
        clear_btn = QPushButton("清空")
        clear_btn.clicked.connect(self.clear_all)
        button_layout.addWidget(clear_btn)
        
        # 保存按钮
        save_btn = QPushButton("保存结果")
        save_btn.clicked.connect(self.save_result)
        button_layout.addWidget(save_btn)
        
        return button_layout

    def add_symptom(self, symptom):
        """添加常见症状到输入框"""
        current_text = self.symptoms_text.toPlainText()
        if current_text:
            self.symptoms_text.setPlainText(f"{current_text}\n{symptom}")
        else:
            self.symptoms_text.setPlainText(symptom)

    def analyze_symptoms(self):
        """分析症状"""
        try:
            # 显示进度条
            self.progress_bar.setVisible(True)
            self.progress_bar.setValue(20)
            self.analyze_btn.setEnabled(False)
            QApplication.processEvents()
            
            # 获取患者信息
            user_info = {
                "age": self.age_input.value(),
                "gender": self.gender_combo.currentText(),
                "height": self.height_input.value(),
                "weight": self.weight_input.value()
            }
            
            symptoms = self.symptoms_text.toPlainText()
            if not symptoms:
                QMessageBox.warning(self, "警告", "请输入症状描述")
                return
            
            # 更新进度
            self.progress_bar.setValue(40)
            QApplication.processEvents()
            
            # 根据选择的模型进行分析
            selected_model = self.model_combo.currentText()
            if selected_model == "OpenAI":
                result = self.analyzer.get_openai_analysis(
                    self.analyzer.build_medical_prompt(user_info, symptoms)
                )
            elif selected_model == "DeepSeek":
                result = self.analyzer.get_deepseek_analysis(
                    self.analyzer.build_medical_prompt(user_info, symptoms)
                )
            else:  # 双模型分析
                result = self.analyzer.analyze(user_info, symptoms)
            
            # 显示结果
            self.output_text.setPlainText(result)
            
            # 完成进度
            self.progress_bar.setValue(100)
            self.statusBar.showMessage("分析完成")
            
        except Exception as e:
            QMessageBox.critical(self, "错误", f"分析失败: {str(e)}")
        finally:
            self.analyze_btn.setEnabled(True)
            self.progress_bar.setVisible(False)

    def clear_all(self):
        """清空所有输入和输出"""
        self.age_input.setValue(0)
        self.height_input.setValue(0)
        self.weight_input.setValue(0)
        self.symptoms_text.clear()
        self.output_text.clear()

    def save_result(self):
        """保存分析结果"""
        if not self.output_text.toPlainText():
            QMessageBox.warning(self, "警告", "没有可保存的结果")
            return
            
        file_name, _ = QFileDialog.getSaveFileName(
            self,
            "保存分析结果",
            "",
            "Text Files (*.txt);;All Files (*)"
        )
        
        if file_name:
            try:
                with open(file_name, 'w', encoding='utf-8') as f:
                    f.write(self.output_text.toPlainText())
                QMessageBox.information(self, "成功", "结果已保存")
            except Exception as e:
                QMessageBox.critical(self, "错误", f"保存失败: {str(e)}")

    def create_menu(self):
        """创建菜单栏"""
        menubar = self.menuBar()
        
        # 文件菜单
        file_menu = menubar.addMenu("文件")
        
        # 新建病历
        new_record = QAction("新建病历", self)
        new_record.setShortcut("Ctrl+N")
        new_record.triggered.connect(self.create_new_record)
        file_menu.addAction(new_record)
        
        # 加载病历
        load_record = QAction("加载病历", self)
        load_record.setShortcut("Ctrl+O")
        load_record.triggered.connect(self.load_medical_record)
        file_menu.addAction(load_record)
        
        file_menu.addSeparator()
        
        # 退出
        exit_action = QAction("退出", self)
        exit_action.setShortcut("Ctrl+Q")
        exit_action.triggered.connect(self.close)
        file_menu.addAction(exit_action)

    def create_new_record(self):
        """新建病历"""
        try:
            # 清空所有输入
            self.clear_all()
            
            # 创建新建病历对话框
            dialog = QDialog(self)
            dialog.setWindowTitle("新建病历")
            dialog.setMinimumWidth(600)
            layout = QVBoxLayout()
            
            # 患者基本信息
            info_group = QGroupBox("患者基本信息")
            info_layout = QFormLayout()
            
            # 姓名
            name_input = QLineEdit()
            info_layout.addRow("姓名:", name_input)
            
            # 身份证号
            id_input = QLineEdit()
            info_layout.addRow("身份证号:", id_input)
            
            # 联系电话
            phone_input = QLineEdit()
            info_layout.addRow("联系电话:", phone_input)
            
            # 年龄
            age_input = QSpinBox()
            age_input.setRange(0, 120)
            info_layout.addRow("年龄:", age_input)
            
            # 性别
            gender_combo = QComboBox()
            gender_combo.addItems(["男", "女"])
            info_layout.addRow("性别:", gender_combo)
            
            # 身高体重
            height_input = QSpinBox()
            height_input.setRange(0, 250)
            height_input.setSuffix(" cm")
            info_layout.addRow("身高:", height_input)
            
            weight_input = QDoubleSpinBox()
            weight_input.setRange(0, 200)
            weight_input.setSuffix(" kg")
            info_layout.addRow("体重:", weight_input)
            
            info_group.setLayout(info_layout)
            layout.addWidget(info_group)
            
            # 就诊信息
            visit_group = QGroupBox("就诊信息")
            visit_layout = QFormLayout()
            
            # 就诊类型
            visit_type = QComboBox()
            visit_type.addItems(["初诊", "复诊", "急诊"])
            visit_layout.addRow("就诊类型:", visit_type)
            
            # 科室
            department = QComboBox()
            department.addItems(["内科", "外科", "儿科", "妇科", "急诊科"])
            visit_layout.addRow("就诊科室:", department)
            
            visit_group.setLayout(visit_layout)
            layout.addWidget(visit_group)
            
            # 按钮
            button_layout = QHBoxLayout()
            save_btn = QPushButton("保存")
            cancel_btn = QPushButton("取消")
            
            def save_new_record():
                try:
                    # 收集患者信息
                    patient_info = {
                        "name": name_input.text(),
                        "id_number": id_input.text(),
                        "phone": phone_input.text(),
                        "age": age_input.value(),
                        "gender": gender_combo.currentText(),
                        "height": height_input.value(),
                        "weight": weight_input.value()
                    }
                    
                    # 更新界面显示
                    self.age_input.setValue(age_input.value())
                    self.gender_combo.setCurrentText(gender_combo.currentText())
                    self.height_input.setValue(height_input.value())
                    self.weight_input.setValue(weight_input.value())
                    
                    # 保存到数据库
                    self.cursor.execute('''
                        INSERT INTO medical_records (
                            timestamp, patient_info, symptoms, diagnosis
                        ) VALUES (?, ?, ?, ?)
                    ''', (
                        datetime.now().strftime('%Y-%m-%d %H:%M:%S'),
                        json.dumps(patient_info),
                        "",  # 空症状
                        ""   # 空诊断
                    ))
                    
                    self.db.commit()
                    
                    QMessageBox.information(dialog, "成功", "新病历已创建")
                    dialog.accept()
                    
                except Exception as e:
                    QMessageBox.warning(dialog, "错误", f"保存病历失败: {str(e)}")
            
            save_btn.clicked.connect(save_new_record)
            cancel_btn.clicked.connect(dialog.reject)
            
            button_layout.addWidget(save_btn)
            button_layout.addWidget(cancel_btn)
            layout.addLayout(button_layout)
            
            dialog.setLayout(layout)
            dialog.exec()
            
        except Exception as e:
            QMessageBox.critical(self, "错误", f"创建病历失败: {str(e)}")

    def save_medical_record(self):
        """保存病历记录到数据库"""
        record = {
            'timestamp': datetime.now().strftime('%Y-%m-%d %H:%M:%S'),
            'patient_info': {
                'age': self.age_input.value(),
                'gender': self.gender_combo.currentText(),
                'height': self.height_input.value(),
                'weight': self.weight_input.value()
            },
            'symptoms': self.symptoms_text.toPlainText(),
            'diagnosis': self.output_text.toPlainText()
        }
        
        self.cursor.execute('''
            INSERT INTO medical_records (timestamp, patient_info, symptoms, diagnosis)
            VALUES (?, ?, ?, ?)
        ''', (
            record['timestamp'],
            json.dumps(record['patient_info']),
            record['symptoms'],
            record['diagnosis']
        ))
        self.db.commit()
        
        # 添加体重记录到健康趋势
        self.add_health_trend('体重', record['patient_info']['weight'])
        
        QMessageBox.information(self, "成功", "病历记录已保存")

    def add_health_trend(self, trend_type: str, value: float):
        """添加健康趋势数据"""
        self.cursor.execute('''
            INSERT INTO health_trends (timestamp, type, value)
            VALUES (?, ?, ?)
        ''', (
            datetime.now().strftime('%Y-%m-%d %H:%M:%S'),
            trend_type,
            value
        ))
        self.db.commit()

    def create_health_trends(self):
        """创建健康趋势图表"""
        from matplotlib.backends.backend_qt5agg import FigureCanvasQTAgg as FigureCanvas
        from matplotlib.figure import Figure
        import matplotlib.dates as mdates
        
        trends_group = QGroupBox("健康趋势")
        layout = QVBoxLayout()
        
        # 创建图表
        self.figure = Figure(figsize=(6, 4))
        self.canvas = FigureCanvas(self.figure)
        self.ax = self.figure.add_subplot(111)
        
        # 添加图表类型选择
        type_layout = QHBoxLayout()
        type_layout.addWidget(QLabel("选择趋势类型:"))
        
        self.trend_type = QComboBox()
        self.trend_type.addItems(['体重', '血压', '血糖', '体温'])
        self.trend_type.currentTextChanged.connect(self.update_health_chart)
        type_layout.addWidget(self.trend_type)
        
        # 添加数据输入
        input_layout = QHBoxLayout()
        self.value_input = QDoubleSpinBox()
        self.value_input.setRange(0, 1000)
        self.value_input.setDecimals(1)
        add_btn = QPushButton("添加数据")
        add_btn.clicked.connect(self.add_trend_data)
        
        input_layout.addWidget(QLabel("数值:"))
        input_layout.addWidget(self.value_input)
        input_layout.addWidget(add_btn)
        
        # 添加时间范围选择
        range_layout = QHBoxLayout()
        self.range_combo = QComboBox()
        self.range_combo.addItems(['最近7天', '最近30天', '最近90天', '全部'])
        self.range_combo.currentTextChanged.connect(self.update_health_chart)
        
        range_layout.addWidget(QLabel("时间范围:"))
        range_layout.addWidget(self.range_combo)
        
        # 添加到主布局
        layout.addLayout(type_layout)
        layout.addLayout(input_layout)
        layout.addLayout(range_layout)
        layout.addWidget(self.canvas)
        
        trends_group.setLayout(layout)
        
        # 初始更新图表
        self.update_health_chart()
        
        return trends_group

    def add_trend_data(self):
        """添加健康趋势数据"""
        try:
            trend_type = self.trend_type.currentText()
            value = self.value_input.value()
            
            if value <= 0:
                QMessageBox.warning(self, "警告", "请输入有效的数值")
                return
            
            # 保存到数据库
            self.cursor.execute('''
                INSERT INTO health_trends (timestamp, type, value)
                VALUES (?, ?, ?)
            ''', (
                datetime.now().strftime('%Y-%m-%d %H:%M:%S'),
                trend_type,
                value
            ))
            self.db.commit()
            
            # 更新图表
            self.update_health_chart()
            
            # 清空输入
            self.value_input.setValue(0)
            
            QMessageBox.information(self, "成功", f"已添加{trend_type}数据")
            
        except Exception as e:
            QMessageBox.warning(self, "警告", f"添加数据失败: {str(e)}")

    def create_symptom_tags(self):
        tag_group = QGroupBox("常见症状标签")
        layout = QVBoxLayout()
        
        # 症状分类
        categories = {
            "全身症状": ["发热", "乏力", "体重减轻", "盗汗"],
            "消化系统": ["腹痛", "恶心", "呕吐", "腹泻"],
            "呼吸系统": ["咳嗽", "咳痰", "胸痛", "呼吸困难"],
            "心血管系统": ["心悸", "胸闷", "水肿"],
            "神经系统": ["头痛", "头晕", "失眠", "抑郁"]
        }
        
        for category, symptoms in categories.items():
            category_layout = QHBoxLayout()
            category_layout.addWidget(QLabel(category))
            for symptom in symptoms:
                btn = QPushButton(symptom)
                btn.setCheckable(True)
                btn.toggled.connect(lambda checked, s=symptom: self.toggle_symptom(s, checked))
                category_layout.addWidget(btn)
            layout.addLayout(category_layout)
        
        tag_group.setLayout(layout)
        return tag_group

    def toggle_symptom(self, symptom: str, checked: bool):
        """切换症状标签"""
        current_text = self.symptoms_text.toPlainText()
        if checked:
            if current_text:
                self.symptoms_text.append(f"\n{symptom}")
            else:
                self.symptoms_text.setText(symptom)
        else:
            # 移除症状
            self.symptoms_text.setText(current_text.replace(symptom, "").strip())

    def create_medication_reminder(self):
        reminder_group = QGroupBox("用药提醒")
        layout = QVBoxLayout()
        
        self.medication_list = QTableWidget()
        self.medication_list.setColumnCount(4)
        self.medication_list.setHorizontalHeaderLabels(['药品名称', '用法用量', '服用时间', '注意事项'])
        self.medication_list.setSelectionMode(QTableWidget.SelectionMode.MultiSelection)  # 允许多选
        
        # 添加按钮布局
        button_layout = QHBoxLayout()
        
        # 原有按钮
        add_btn = QPushButton("手动添加")
        extract_btn = QPushButton("从分析结果提取")
        edit_btn = QPushButton("编辑")
        delete_btn = QPushButton("删除")
        
        # 新增按钮
        batch_edit_btn = QPushButton("批量编辑")
        schedule_btn = QPushButton("生成时间表")
        check_interaction_btn = QPushButton("相互作用检查")
        export_btn = QPushButton("导出提醒")
        
        # 连接信号
        add_btn.clicked.connect(self.add_medication_reminder)
        extract_btn.clicked.connect(self.extract_medication_from_analysis)
        edit_btn.clicked.connect(self.edit_medication_reminder)
        delete_btn.clicked.connect(self.delete_medication_reminder)
        batch_edit_btn.clicked.connect(self.batch_edit_medications)
        schedule_btn.clicked.connect(self.generate_medication_schedule)
        check_interaction_btn.clicked.connect(self.check_drug_interactions)
        export_btn.clicked.connect(self.export_medication_reminders)
        
        # 添加按钮到布局
        for btn in [add_btn, extract_btn, edit_btn, delete_btn, 
                    batch_edit_btn, schedule_btn, check_interaction_btn, export_btn]:
            button_layout.addWidget(btn)
        
        layout.addWidget(self.medication_list)
        layout.addLayout(button_layout)
        reminder_group.setLayout(layout)
        return reminder_group

    def create_medication_dialog(self):
        """创建用药信息输入对话框"""
        dialog = QDialog(self)
        dialog.setWindowTitle("添加用药提醒")
        layout = QVBoxLayout()

        # 药品信息输入表单
        form_layout = QFormLayout()
        
        # 药品名称
        medicine_name = QLineEdit()
        form_layout.addRow("药品名称:", medicine_name)
        
        # 用法用量
        dosage = QLineEdit()
        dosage.setPlaceholderText("如: 5mg/片")
        form_layout.addRow("用法用量:", dosage)
        
        # 服用时间
        time_combo = QComboBox()
        time_combo.addItems([
            "每日一次(早餐后)",
            "每日两次(早晚餐后)",
            "每日三次(三餐后)",
            "每日四次(每6小时)",
            "需要时服用",
            "自定义..."
        ])
        form_layout.addRow("服用时间:", time_combo)
        
        # 自定义时间输入
        custom_time = QLineEdit()
        custom_time.setPlaceholderText("自定义服用时间")
        custom_time.setVisible(False)
        form_layout.addRow("自定义:", custom_time)
        
        # 服用天数
        days_spin = QSpinBox()
        days_spin.setRange(1, 365)
        days_spin.setValue(7)
        days_spin.setSuffix(" 天")
        form_layout.addRow("服用天数:", days_spin)
        
        # 注意事项
        notes = QTextEdit()
        notes.setPlaceholderText("输入用药注意事项...")
        notes.setMaximumHeight(100)
        form_layout.addRow("注意事项:", notes)
        
        layout.addLayout(form_layout)
        
        # 按钮
        button_layout = QHBoxLayout()
        save_btn = QPushButton("保存")
        cancel_btn = QPushButton("取消")
        
        button_layout.addWidget(save_btn)
        button_layout.addWidget(cancel_btn)
        layout.addLayout(button_layout)
        
        dialog.setLayout(layout)
        
        # 连接信号
        def on_time_changed(text):
            custom_time.setVisible(text == "自定义...")
            
        time_combo.currentTextChanged.connect(on_time_changed)
        save_btn.clicked.connect(dialog.accept)
        cancel_btn.clicked.connect(dialog.reject)
        
        # 返回对话框和输入控件
        return dialog, {
            'name': medicine_name,
            'dosage': dosage,
            'time': time_combo,
            'custom_time': custom_time,
            'days': days_spin,
            'notes': notes
        }

    def add_medication_reminder(self):
        """添加用药提醒"""
        try:
            # 创建输入对话框
            dialog, inputs = self.create_medication_dialog()
            
            if dialog.exec() == QDialog.DialogCode.Accepted:
                # 获取输入的值
                medicine_name = inputs['name'].text()
                dosage = inputs['dosage'].text()
                time = inputs['time'].currentText()
                if time == "自定义...":
                    time = inputs['custom_time'].text()
                days = inputs['days'].value()
                notes = inputs['notes'].toPlainText()
                
                if not medicine_name or not dosage:
                    QMessageBox.warning(self, "警告", "药品名称和用法用量不能为空")
                    return
                
                # 添加到用药提醒表格
                row = self.medication_list.rowCount()
                self.medication_list.insertRow(row)
                self.medication_list.setItem(row, 0, QTableWidgetItem(medicine_name))
                self.medication_list.setItem(row, 1, QTableWidgetItem(dosage))
                self.medication_list.setItem(row, 2, QTableWidgetItem(time))
                self.medication_list.setItem(row, 3, QTableWidgetItem(notes))
                
                # 保存到数据库
                self.cursor.execute('''
                    INSERT INTO medication_reminders 
                    (medicine_name, dosage, time, notes)
                    VALUES (?, ?, ?, ?)
                ''', (medicine_name, dosage, time, notes))
                
                self.db.commit()
                
                # 添加到系统提醒（如果需要）
                self.schedule_medication_reminder(medicine_name, time, days)
                
                QMessageBox.information(self, "成功", "已添加用药提醒")
                
        except Exception as e:
            QMessageBox.warning(self, "警告", f"添加用药提醒失败: {str(e)}")

    def schedule_medication_reminder(self, medicine_name: str, time_str: str, days: int):
        """设置系统用药提醒"""
        try:
            from datetime import datetime, timedelta
            
            # 根据不同操作系统设置提醒
            system = platform.system()
            
            if system == "Darwin":  # macOS
                import os
                for day in range(days):
                    cmd = f'''osascript -e 'display notification "请按时服用 {medicine_name}" with title "用药提醒" sound name "Glass"' '''
                    # 使用 launchctl 设置定时任务
                    plist_content = f'''<?xml version="1.0" encoding="UTF-8"?>
                    <!DOCTYPE plist PUBLIC "-//Apple//DTD PLIST 1.0//EN" "http://www.apple.com/DTDs/PropertyList-1.0.dtd">
                    <plist version="1.0">
                    <dict>
                        <key>Label</key>
                        <string>com.medical.reminder.{medicine_name.replace(" ", "")}</string>
                        <key>ProgramArguments</key>
                        <array>
                            <string>/usr/bin/osascript</string>
                            <string>-e</string>
                            <string>display notification "请按时服用 {medicine_name}" with title "用药提醒" sound name "Glass"</string>
                        </array>
                        <key>StartCalendarInterval</key>
                        <dict>
                            <key>Hour</key>
                            <integer>9</integer>
                            <key>Minute</key>
                            <integer>0</integer>
                        </dict>
                    </dict>
                    </plist>'''
                    
                    # 保存 plist 文件
                    plist_path = os.path.expanduser(f"~/Library/LaunchAgents/com.medical.reminder.{medicine_name.replace(' ', '')}.plist")
                    with open(plist_path, 'w') as f:
                        f.write(plist_content)
                    
                    # 加载 plist
                    os.system(f"launchctl load {plist_path}")
                    
            elif system == "Windows":
                import win32com.client
                scheduler = win32com.client.Dispatch('Schedule.Service')
                scheduler.Connect()
                root_folder = scheduler.GetFolder('\\')
                task_def = scheduler.NewTask(0)
                
                # 创建触发器
                start_time = datetime.now()
                end_time = start_time + timedelta(days=days)
                
                TASK_TRIGGER_TIME = 1
                trigger = task_def.Triggers.Create(TASK_TRIGGER_TIME)
                trigger.StartBoundary = start_time.isoformat()
                trigger.EndBoundary = end_time.isoformat()
                
                # 创建操作
                TASK_ACTION_EXEC = 0
                action = task_def.Actions.Create(TASK_ACTION_EXEC)
                action.Path = "cmd.exe"
                action.Arguments = f'/c msg * "请按时服用 {medicine_name}"'
                
                # 注册任务
                task_name = f"MedicalReminder_{medicine_name.replace(' ', '_')}"
                root_folder.RegisterTaskDefinition(
                    task_name,
                    task_def,
                    6,  # TASK_CREATE_OR_UPDATE
                    None,  # 用户名
                    None,  # 密码
                    0  # TASK_LOGON_NONE
                )
                
        except Exception as e:
            print(f"设置系统提醒失败: {str(e)}")

    def create_smart_inquiry(self):
        inquiry_group = QGroupBox("智能问诊")
        layout = QVBoxLayout()
        
        # 主诉引导
        chief_complaint = QComboBox()
        chief_complaint.addItems(['头痛', '腹痛', '发热', '咳嗽'])
        chief_complaint.currentTextChanged.connect(self.update_inquiry_questions)
        
        # 动态问题列表
        self.question_list = QListWidget()
        
        # 答案输入区
        self.answer_input = QTextEdit()
        
        layout.addWidget(QLabel("主要症状:"))
        layout.addWidget(chief_complaint)
        layout.addWidget(self.question_list)
        layout.addWidget(self.answer_input)
        
        inquiry_group.setLayout(layout)
        return inquiry_group

    def update_inquiry_questions(self, symptom: str):
        """根据主诉更新问诊问题"""
        questions = {
            "头痛": [
                "疼痛性质？(胀痛/刺痛/钝痛)",
                "持续时间？",
                "是否伴有其他症状？",
                "有无诱因？"
            ],
            # 其他症状的问题...
        }
        
        self.question_list.clear()
        self.question_list.addItems(questions.get(symptom, []))

    def load_medical_record(self):
        """加载病历记录"""
        try:
            # 获取最近的病历记录
            self.cursor.execute('''
                SELECT id, timestamp, patient_info, symptoms, diagnosis 
                FROM medical_records 
                ORDER BY timestamp DESC
            ''')
            
            records = self.cursor.fetchall()
            if not records:
                QMessageBox.information(self, "提示", "没有找到病历记录")
                return
            
            # 创建选择对话框
            dialog = QDialog(self)
            dialog.setWindowTitle("病历记录管理")
            dialog.setMinimumWidth(600)
            layout = QVBoxLayout()
            
            # 创建表格
            table = QTableWidget()
            table.setColumnCount(4)
            table.setHorizontalHeaderLabels(['时间', '患者信息', '症状', '诊断'])
            table.horizontalHeader().setSectionResizeMode(QHeaderView.ResizeMode.ResizeToContents)
            
            # 填充数据
            for record in records:
                row = table.rowCount()
                table.insertRow(row)
                
                # 时间
                table.setItem(row, 0, QTableWidgetItem(record[1]))
                
                # 患者信息
                patient_info = json.loads(record[2])
                info_text = f"{patient_info['age']}岁{patient_info['gender']}"
                table.setItem(row, 1, QTableWidgetItem(info_text))
                
                # 症状（显示前50个字符）
                symptoms = record[3][:50] + "..." if len(record[3]) > 50 else record[3]
                table.setItem(row, 2, QTableWidgetItem(symptoms))
                
                # 诊断（显示前50个字符）
                diagnosis = record[4][:50] + "..." if len(record[4]) > 50 else record[4]
                table.setItem(row, 3, QTableWidgetItem(diagnosis))
            
            layout.addWidget(table)
            
            # 按钮布局
            buttons = QHBoxLayout()
            load_btn = QPushButton("加载")
            delete_btn = QPushButton("删除")
            cancel_btn = QPushButton("取消")
            
            buttons.addWidget(load_btn)
            buttons.addWidget(delete_btn)
            buttons.addWidget(cancel_btn)
            layout.addLayout(buttons)
            
            dialog.setLayout(layout)
            
            # 删除病历记录
            def delete_record():
                current_row = table.currentRow()
                if current_row < 0:
                    QMessageBox.warning(dialog, "警告", "请先选择要删除的记录")
                    return
                
                if QMessageBox.question(
                    dialog,
                    "确认删除",
                    "确定要删除这条病历记录吗？此操作不可恢复。",
                    QMessageBox.StandardButton.Yes | QMessageBox.StandardButton.No
                ) == QMessageBox.StandardButton.Yes:
                    try:
                        record_id = records[current_row][0]
                        
                        # 删除相关处方
                        self.cursor.execute('''
                            DELETE FROM prescription_items 
                            WHERE prescription_id IN (
                                SELECT id FROM prescriptions 
                                WHERE patient_name = (
                                    SELECT json_extract(patient_info, '$.name') 
                                    FROM medical_records 
                                    WHERE id = ?
                                )
                            )
                        ''', (record_id,))
                        
                        self.cursor.execute('''
                            DELETE FROM prescriptions 
                            WHERE patient_name = (
                                SELECT json_extract(patient_info, '$.name') 
                                FROM medical_records 
                                WHERE id = ?
                            )
                        ''', (record_id,))
                        
                        # 删除病历记录
                        self.cursor.execute('DELETE FROM medical_records WHERE id = ?', (record_id,))
                        self.db.commit()
                        
                        # 从表格中移除
                        table.removeRow(current_row)
                        records.pop(current_row)
                        
                        # 更新处方列表
                        self.update_prescription_list()
                        
                        QMessageBox.information(dialog, "成功", "病历记录已删除")
                        
                    except Exception as e:
                        QMessageBox.warning(dialog, "错误", f"删除失败: {str(e)}")
            
            # 加载病历记录
            def load_record():
                current_row = table.currentRow()
                if current_row >= 0:
                    selected = records[current_row]
                    patient_info = json.loads(selected[2])
                    
                    # 填充数据
                    self.age_input.setValue(patient_info['age'])
                    self.gender_combo.setCurrentText(patient_info['gender'])
                    self.height_input.setValue(patient_info.get('height', 0))
                    self.weight_input.setValue(patient_info.get('weight', 0))
                    self.symptoms_text.setPlainText(selected[3])
                    self.output_text.setPlainText(selected[4])
                    
                    # 更新处方列表
                    self.update_prescription_list()
                    
                    QMessageBox.information(dialog, "成功", "病历记录已加载")
                    dialog.accept()
                else:
                    QMessageBox.warning(dialog, "警告", "请先选择要加载的记录")
            
            # 连接信号
            load_btn.clicked.connect(load_record)
            delete_btn.clicked.connect(delete_record)
            cancel_btn.clicked.connect(dialog.reject)
            
            # 双击加载记录
            table.cellDoubleClicked.connect(lambda: load_record())
            
            dialog.exec()
            
        except Exception as e:
            QMessageBox.critical(self, "错误", f"加载病历失败: {str(e)}")

    def export_medical_record(self):
        """导出病历记录"""
        try:
            file_name, _ = QFileDialog.getSaveFileName(
                self,
                "导出病历",
                "",
                "Word文档 (*.docx);;PDF文件 (*.pdf);;文本文件 (*.txt)"
            )
            
            if not file_name:
                return
            
            if file_name.endswith('.txt'):
                # 导出为文本文件
                content = f"""病历记录

时间: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}

患者信息:
- 年龄: {self.age_input.value()}岁
- 性别: {self.gender_combo.currentText()}
- 身高: {self.height_input.value()}cm
- 体重: {self.weight_input.value()}kg

症状描述:
{self.symptoms_text.toPlainText()}

诊断结果:
{self.output_text.toPlainText()}
"""
                with open(file_name, 'w', encoding='utf-8') as f:
                    f.write(content)
                
            elif file_name.endswith('.docx'):
                from docx import Document
                doc = Document()
                
                # 添加标题
                doc.add_heading('病历记录', 0)
                
                # 添加时间
                doc.add_paragraph(f"时间: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
                
                # 添加患者信息
                doc.add_heading('患者信息', level=1)
                info = doc.add_paragraph()
                info.add_run(f"年龄: {self.age_input.value()}岁\n")
                info.add_run(f"性别: {self.gender_combo.currentText()}\n")
                info.add_run(f"身高: {self.height_input.value()}cm\n")
                info.add_run(f"体重: {self.weight_input.value()}kg")
                
                # 添加症状描述
                doc.add_heading('症状描述', level=1)
                doc.add_paragraph(self.symptoms_text.toPlainText())
                
                # 添加诊断结果
                doc.add_heading('诊断结果', level=1)
                doc.add_paragraph(self.output_text.toPlainText())
                
                doc.save(file_name)
            
            QMessageBox.information(self, "成功", "病历记录已导出")
            
        except Exception as e:
            QMessageBox.critical(self, "错误", f"导出病历失败: {str(e)}")

    def extract_medication_from_analysis(self):
        """从分析结果中提取用药信息"""
        try:
            # 获取选中的文本，如果没有选中则使用全部文本
            cursor = self.output_text.textCursor()
            selected_text = cursor.selectedText()
            
            # 创建选择对话框
            dialog = QDialog(self)
            dialog.setWindowTitle("选择要提取的文本段落")
            dialog.setMinimumWidth(600)
            layout = QVBoxLayout()
            
            # 说明标签
            label = QLabel("请选择要提取用药信息的文本段落：")
            layout.addWidget(label)
            
            # 文本编辑框
            text_edit = QTextEdit()
            text_edit.setPlainText(selected_text if selected_text else self.output_text.toPlainText())
            text_edit.setMinimumHeight(200)
            layout.addWidget(text_edit)
            
            # 快速选择按钮
            quick_select_layout = QHBoxLayout()
            sections = ["用药方案", "治疗方案", "推荐用药", "用药建议"]
            
            def select_section(section_name):
                full_text = self.output_text.toPlainText()
                sections = full_text.split("===")
                for section in sections:
                    if section_name in section:
                        text_edit.setPlainText(section.strip())
                        return
            
            for section in sections:
                btn = QPushButton(section)
                btn.clicked.connect(lambda checked, s=section: select_section(s))
                quick_select_layout.addWidget(btn)
            
            layout.addLayout(quick_select_layout)
            
            # 按钮
            button_layout = QHBoxLayout()
            extract_btn = QPushButton("提取用药信息")
            cancel_btn = QPushButton("取消")
            
            button_layout.addWidget(extract_btn)
            button_layout.addWidget(cancel_btn)
            layout.addLayout(button_layout)
            
            dialog.setLayout(layout)
            
            # 连接信号
            extract_btn.clicked.connect(dialog.accept)
            cancel_btn.clicked.connect(dialog.reject)
            
            # 显示对话框
            if dialog.exec() != QDialog.DialogCode.Accepted:
                return
            
            # 获取选择的文本
            selected_text = text_edit.toPlainText()
            if not selected_text:
                QMessageBox.warning(self, "警告", "请选择要提取的文本")
                return
            
            print("Trying to match text:", selected_text)  # 调试输出
            
            # 提取用药信息
            medications = []
            patterns = [
                # 模式1：标准格式（匹配带缩进的完整格式）
                r"""
                    ^-\s*([^：\n]+)：  # 药品名称
                    \s*([^，\n]+?)  # 剂量（更宽松的匹配）
                    (?:，|\s+)([^\n]+)  # 用法（允许逗号分隔）
                    \s+用药说明：([^\n]+)  # 用药说明
                    \s+注意事项：([^\n]+)  # 注意事项
                """,
                
                # 模式2：简单格式（只匹配基本信息）
                r"-\s*([^：\n]+)：\s*([^，\n]+?)(?:，|\s+)([^\n]+)",
                
                # 模式3：备用格式（更宽松的匹配）
                r"^-\s*([^：\n]+)：\s*([^，\n]+?)(?:，|\s+)([^\n]+)"
            ]
            
            for i, pattern in enumerate(patterns, 1):
                matches = re.finditer(pattern, selected_text, re.MULTILINE | re.VERBOSE)
                for match in matches:
                    print(f"Pattern {i} found match:", match.groups())  # 调试输出
                    
                    # 跳过非药品行
                    if "无抗生素推荐" in match.group(1):
                        continue
                        
                    med = {
                        'name': match.group(1).strip(),
                        'dosage': match.group(2).strip(),
                        'time': match.group(3).strip(),
                    }
                    
                    # 如果匹配到了说明和注意事项
                    if len(match.groups()) >= 5:
                        notes = []
                        if match.group(4):
                            notes.append(f"说明：{match.group(4).strip()}")
                        if match.group(5):
                            notes.append(f"注意：{match.group(5).strip()}")
                        med['notes'] = "\n".join(notes) if notes else ""
                    else:
                        med['notes'] = ""
                    
                    # 检查是否已存在相同药品
                    if not any(m['name'] == med['name'] for m in medications):
                        medications.append(med)
                        print(f"Added medication: {med}")  # 调试输出
            
            if not medications:
                QMessageBox.information(self, "提示", "在选中文本中未找到用药信息")
                return
            
            # 显示提取结果供确认
            dialog = QDialog(self)
            dialog.setWindowTitle("确认用药信息")
            layout = QVBoxLayout()
            
            # 创建表格显示提取的用药信息
            table = QTableWidget()
            table.setColumnCount(4)
            table.setHorizontalHeaderLabels(['药品名称', '用法用量', '服用时间', '注意事项'])
            table.setRowCount(len(medications))
            
            # 设置表格列宽
            table.horizontalHeader().setSectionResizeMode(0, QHeaderView.ResizeMode.ResizeToContents)
            table.horizontalHeader().setSectionResizeMode(1, QHeaderView.ResizeMode.ResizeToContents)
            table.horizontalHeader().setSectionResizeMode(2, QHeaderView.ResizeMode.ResizeToContents)
            table.horizontalHeader().setSectionResizeMode(3, QHeaderView.ResizeMode.Stretch)
            
            for i, med in enumerate(medications):
                table.setItem(i, 0, QTableWidgetItem(med['name']))
                table.setItem(i, 1, QTableWidgetItem(med['dosage']))
                table.setItem(i, 2, QTableWidgetItem(med['time']))
                table.setItem(i, 3, QTableWidgetItem(med['notes']))
            
            layout.addWidget(table)
            
            # 添加按钮
            button_layout = QHBoxLayout()
            confirm_btn = QPushButton("确认添加")
            cancel_btn = QPushButton("取消")
            edit_btn = QPushButton("编辑")
            
            def edit_selected():
                current_row = table.currentRow()
                if current_row >= 0:
                    dialog, inputs = self.create_medication_dialog()
                    med = medications[current_row]
                    
                    # 填充现有数据
                    inputs['name'].setText(med['name'])
                    inputs['dosage'].setText(med['dosage'])
                    inputs['time'].setText(med['time'])
                    inputs['notes'].setText(med['notes'])
                    
                    if dialog.exec() == QDialog.DialogCode.Accepted:
                        # 更新表格和数据
                        med['name'] = inputs['name'].text()
                        med['dosage'] = inputs['dosage'].text()
                        med['time'] = inputs['time'].currentText()
                        med['notes'] = inputs['notes'].toPlainText()
                        
                        # 更新表格显示
                        table.setItem(current_row, 0, QTableWidgetItem(med['name']))
                        table.setItem(current_row, 1, QTableWidgetItem(med['dosage']))
                        table.setItem(current_row, 2, QTableWidgetItem(med['time']))
                        table.setItem(current_row, 3, QTableWidgetItem(med['notes']))
            
            edit_btn.clicked.connect(edit_selected)
            confirm_btn.clicked.connect(dialog.accept)
            cancel_btn.clicked.connect(dialog.reject)
            
            button_layout.addWidget(edit_btn)
            button_layout.addWidget(confirm_btn)
            button_layout.addWidget(cancel_btn)
            layout.addLayout(button_layout)
            
            dialog.setLayout(layout)
            
            if dialog.exec() == QDialog.DialogCode.Accepted:
                # 添加确认的用药信息
                for med in medications:
                    row = self.medication_list.rowCount()
                    self.medication_list.insertRow(row)
                    self.medication_list.setItem(row, 0, QTableWidgetItem(med['name']))
                    self.medication_list.setItem(row, 1, QTableWidgetItem(med['dosage']))
                    self.medication_list.setItem(row, 2, QTableWidgetItem(med['time']))
                    self.medication_list.setItem(row, 3, QTableWidgetItem(med['notes']))
                    
                    # 保存到数据库
                    self.cursor.execute('''
                        INSERT INTO medication_reminders 
                        (medicine_name, dosage, time, notes)
                        VALUES (?, ?, ?, ?)
                    ''', (med['name'], med['dosage'], med['time'], med['notes']))
                
                self.db.commit()
                QMessageBox.information(self, "成功", f"已添加 {len(medications)} 条用药提醒")
                
        except Exception as e:
            QMessageBox.warning(self, "警告", f"提取用药信息失败: {str(e)}")

    def edit_medication_reminder(self):
        """编辑用药提醒"""
        current_row = self.medication_list.currentRow()
        if current_row < 0:
            QMessageBox.warning(self, "警告", "请先选择要编辑的用药提醒")
            return
        
        try:
            # 获取当前选中的用药信息
            medicine_name = self.medication_list.item(current_row, 0).text()
            dosage = self.medication_list.item(current_row, 1).text()
            time = self.medication_list.item(current_row, 2).text()
            notes = self.medication_list.item(current_row, 3).text()
            
            # 创建编辑对话框
            dialog, inputs = self.create_medication_dialog()
            
            # 填充现有数据
            inputs['name'].setText(medicine_name)
            inputs['dosage'].setText(dosage)
            inputs['time'].setCurrentText(time)
            inputs['notes'].setText(notes)
            
            if dialog.exec() == QDialog.DialogCode.Accepted:
                # 更新表格
                self.medication_list.setItem(current_row, 0, QTableWidgetItem(inputs['name'].text()))
                self.medication_list.setItem(current_row, 1, QTableWidgetItem(inputs['dosage'].text()))
                self.medication_list.setItem(current_row, 2, QTableWidgetItem(inputs['time'].currentText()))
                self.medication_list.setItem(current_row, 3, QTableWidgetItem(inputs['notes'].toPlainText()))
                
                # 更新数据库
                self.cursor.execute('''
                    UPDATE medication_reminders 
                    SET medicine_name = ?, dosage = ?, time = ?, notes = ?
                    WHERE medicine_name = ? AND dosage = ? AND time = ?
                ''', (
                    inputs['name'].text(),
                    inputs['dosage'].text(),
                    inputs['time'].currentText(),
                    inputs['notes'].toPlainText(),
                    medicine_name,
                    dosage,
                    time
                ))
                
                self.db.commit()
                QMessageBox.information(self, "成功", "用药提醒已更新")
                
        except Exception as e:
            QMessageBox.warning(self, "警告", f"编辑用药提醒失败: {str(e)}")

    def delete_medication_reminder(self):
        """删除用药提醒"""
        current_row = self.medication_list.currentRow()
        if current_row < 0:
            QMessageBox.warning(self, "警告", "请先选择要删除的用药提醒")
            return
        
        try:
            medicine_name = self.medication_list.item(current_row, 0).text()
            if QMessageBox.question(
                self,
                "确认删除",
                f"确定要删除 {medicine_name} 的用药提醒吗？"
            ) == QMessageBox.Yes:
                # 从数据库删除
                self.cursor.execute('''
                    DELETE FROM medication_reminders 
                    WHERE medicine_name = ?
                ''', (medicine_name,))
                
                self.db.commit()
                
                # 从表格删除
                self.medication_list.removeRow(current_row)
                
                QMessageBox.information(self, "成功", "用药提醒已删除")
                
        except Exception as e:
            QMessageBox.warning(self, "警告", f"删除用药提醒失败: {str(e)}")

    def batch_edit_medications(self):
        """批量编辑用药提醒"""
        selected_rows = set(item.row() for item in self.medication_list.selectedItems())
        if not selected_rows:
            QMessageBox.warning(self, "警告", "请先选择要编辑的用药提醒")
            return
        
        try:
            dialog = QDialog(self)
            dialog.setWindowTitle("批量编辑用药提醒")
            layout = QVBoxLayout()
            
            # 创建选项
            form_layout = QFormLayout()
            
            # 服用时间
            time_combo = QComboBox()
            time_combo.addItems([
                "保持原值",
                "每日一次(早餐后)",
                "每日两次(早晚餐后)",
                "每日三次(三餐后)",
                "每日四次(每6小时)"
            ])
            form_layout.addRow("服用时间:", time_combo)
            
            # 注意事项
            notes_edit = QTextEdit()
            notes_edit.setPlaceholderText("如需修改注意事项，请在此输入...")
            form_layout.addRow("注意事项:", notes_edit)
            
            layout.addLayout(form_layout)
            
            # 按钮
            buttons = QHBoxLayout()
            save_btn = QPushButton("保存")
            cancel_btn = QPushButton("取消")
            
            save_btn.clicked.connect(dialog.accept)
            cancel_btn.clicked.connect(dialog.reject)
            
            buttons.addWidget(save_btn)
            buttons.addWidget(cancel_btn)
            layout.addLayout(buttons)
            
            dialog.setLayout(layout)
            
            if dialog.exec() == QDialog.DialogCode.Accepted:
                # 更新选中的行
                for row in selected_rows:
                    if time_combo.currentText() != "保持原值":
                        self.medication_list.setItem(row, 2, QTableWidgetItem(time_combo.currentText()))
                    
                    if notes_edit.toPlainText():
                        self.medication_list.setItem(row, 3, QTableWidgetItem(notes_edit.toPlainText()))
                    
                    # 更新数据库
                    self.cursor.execute('''
                        UPDATE medication_reminders 
                        SET time = ?, notes = ?
                        WHERE medicine_name = ?
                    ''', (
                        time_combo.currentText() if time_combo.currentText() != "保持原值" else self.medication_list.item(row, 2).text(),
                        notes_edit.toPlainText() if notes_edit.toPlainText() else self.medication_list.item(row, 3).text(),
                        self.medication_list.item(row, 0).text()
                    ))
                
                self.db.commit()
                QMessageBox.information(self, "成功", f"已更新 {len(selected_rows)} 条用药提醒")
                
        except Exception as e:
            QMessageBox.warning(self, "警告", f"批量编辑失败: {str(e)}")

    def generate_medication_schedule(self):
        """生成用药时间表"""
        try:
            from datetime import datetime, timedelta
            
            # 获取所有用药信息
            medications = []
            for row in range(self.medication_list.rowCount()):
                medications.append({
                    'name': self.medication_list.item(row, 0).text(),
                    'dosage': self.medication_list.item(row, 1).text(),
                    'time': self.medication_list.item(row, 2).text(),
                    'notes': self.medication_list.item(row, 3).text()
                })
            
            if not medications:
                QMessageBox.warning(self, "警告", "没有用药信息可生成时间表")
                return
            
            # 创建时间表对话框
            dialog = QDialog(self)
            dialog.setWindowTitle("用药时间表")
            dialog.setMinimumWidth(800)
            layout = QVBoxLayout()
            
            # 创建日历视图
            calendar = QCalendarWidget()
            layout.addWidget(calendar)
            
            # 创建时间表
            table = QTableWidget()
            table.setColumnCount(4)
            table.setHorizontalHeaderLabels(['时间', '药品', '用量', '注意事项'])
            
            # 设置表格列宽
            table.horizontalHeader().setSectionResizeMode(0, QHeaderView.ResizeMode.ResizeToContents)
            table.horizontalHeader().setSectionResizeMode(1, QHeaderView.ResizeMode.ResizeToContents)
            table.horizontalHeader().setSectionResizeMode(2, QHeaderView.ResizeMode.ResizeToContents)
            table.horizontalHeader().setSectionResizeMode(3, QHeaderView.ResizeMode.Stretch)
            
            def update_schedule():
                table.setRowCount(0)
                current_row = 0
                
                for med in medications:
                    times = []
                    if "每日一次" in med['time']:
                        times = ["08:00"]
                    elif "每日两次" in med['time']:
                        times = ["08:00", "20:00"]
                    elif "每日三次" in med['time']:
                        times = ["08:00", "14:00", "20:00"]
                    elif "每日四次" in med['time']:
                        times = ["06:00", "12:00", "18:00", "24:00"]
                    else:
                        # 处理自定义时间
                        times = ["08:00"]  # 默认时间
                    
                    for t in times:
                        table.insertRow(current_row)
                        table.setItem(current_row, 0, QTableWidgetItem(t))
                        table.setItem(current_row, 1, QTableWidgetItem(med['name']))
                        table.setItem(current_row, 2, QTableWidgetItem(med['dosage']))
                        table.setItem(current_row, 3, QTableWidgetItem(med['notes']))
                        current_row += 1
                
                # 按时间排序
                table.sortItems(0)
            
            # 初始更新时间表
            update_schedule()
            
            # 连接日历选择信号
            calendar.selectionChanged.connect(update_schedule)
            
            layout.addWidget(table)
            
            # 添加导出和设置按钮
            button_layout = QHBoxLayout()
            
            # 导出按钮
            export_btn = QPushButton("导出时间表")
            def export_schedule():
                file_name, _ = QFileDialog.getSaveFileName(
                    dialog,
                    "导出时间表",
                    "",
                    "Excel文件 (*.xlsx);;CSV文件 (*.csv)"
                )
                
                if file_name:
                    selected_date = calendar.selectedDate().toString("yyyy-MM-dd")
                    if file_name.endswith('.xlsx'):
                        import pandas as pd
                        data = []
                        for row in range(table.rowCount()):
                            data.append({
                                '日期': selected_date,
                                '时间': table.item(row, 0).text(),
                                '药品': table.item(row, 1).text(),
                                '用量': table.item(row, 2).text(),
                                '注意事项': table.item(row, 3).text()
                            })
                        df = pd.DataFrame(data)
                        df.to_excel(file_name, index=False)
                    else:
                        with open(file_name, 'w', encoding='utf-8') as f:
                            f.write('日期,时间,药品,用量,注意事项\n')
                            for row in range(table.rowCount()):
                                f.write(f"{selected_date},{table.item(row, 0).text()},{table.item(row, 1).text()},{table.item(row, 2).text()},{table.item(row, 3).text()}\n")
                    
                    QMessageBox.information(dialog, "成功", "时间表已导出")
            
            export_btn.clicked.connect(export_schedule)
            button_layout.addWidget(export_btn)
            
            # 设置按钮
            settings_btn = QPushButton("时间设置")
            def show_time_settings():
                settings_dialog = QDialog(dialog)
                settings_dialog.setWindowTitle("时间设置")
                settings_layout = QFormLayout()
                
                time_settings = {
                    "每日一次": QTimeEdit(QTime(8, 0)),
                    "每日两次-早": QTimeEdit(QTime(8, 0)),
                    "每日两次-晚": QTimeEdit(QTime(20, 0)),
                    "每日三次-早": QTimeEdit(QTime(8, 0)),
                    "每日三次-午": QTimeEdit(QTime(14, 0)),
                    "每日三次-晚": QTimeEdit(QTime(20, 0))
                }
                
                for label, time_edit in time_settings.items():
                    settings_layout.addRow(label, time_edit)
                
                save_btn = QPushButton("保存")
                save_btn.clicked.connect(settings_dialog.accept)
                settings_layout.addRow(save_btn)
                
                settings_dialog.setLayout(settings_layout)
                
                if settings_dialog.exec() == QDialog.DialogCode.Accepted:
                    update_schedule()  # 使用新的时间设置更新时间表
            
            settings_btn.clicked.connect(show_time_settings)
            button_layout.addWidget(settings_btn)
            
            layout.addLayout(button_layout)
            
            dialog.setLayout(layout)
            dialog.exec()
            
        except Exception as e:
            QMessageBox.warning(self, "警告", f"生成时间表失败: {str(e)}")

    def check_drug_interactions(self):
        """检查药物相互作用"""
        try:
            medications = []
            for row in range(self.medication_list.rowCount()):
                medications.append(self.medication_list.item(row, 0).text())
            
            if len(medications) < 2:
                QMessageBox.information(self, "提示", "需要至少两种药物才能检查相互作用")
                return
            
            # 调用AI分析药物相互作用
            prompt = f"""请分析以下药物之间可能存在的相互作用：
{', '.join(medications)}

请按以下格式输出：
1. 存在的相互作用
2. 风险等级（高/中/低）
3. 注意事项
4. 建议措施"""
            
            result = self.analyzer.get_openai_analysis(prompt)
            
            # 显示结果
            dialog = QDialog(self)
            dialog.setWindowTitle("药物相互作用分析")
            layout = QVBoxLayout()
            
            text_browser = QTextBrowser()
            text_browser.setText(result)
            layout.addWidget(text_browser)
            
            close_btn = QPushButton("关闭")
            close_btn.clicked.connect(dialog.accept)
            layout.addWidget(close_btn)
            
            dialog.setLayout(layout)
            dialog.exec()
            
        except Exception as e:
            QMessageBox.warning(self, "警告", f"检查药物相互作用失败: {str(e)}")

    def export_medication_reminders(self):
        """导出用药提醒"""
        try:
            file_name, _ = QFileDialog.getSaveFileName(
                self,
                "导出用药提醒",
                "",
                "Excel文件 (*.xlsx);;CSV文件 (*.csv);;iCal文件 (*.ics)"
            )
            
            if not file_name:
                return
            
            medications = []
            for row in range(self.medication_list.rowCount()):
                medications.append({
                    'name': self.medication_list.item(row, 0).text(),
                    'dosage': self.medication_list.item(row, 1).text(),
                    'time': self.medication_list.item(row, 2).text(),
                    'notes': self.medication_list.item(row, 3).text()
                })
            
            if file_name.endswith('.xlsx'):
                import pandas as pd
                df = pd.DataFrame(medications)
                df.to_excel(file_name, index=False)
                
            elif file_name.endswith('.csv'):
                import csv
                with open(file_name, 'w', encoding='utf-8', newline='') as f:
                    writer = csv.DictWriter(f, fieldnames=['name', 'dosage', 'time', 'notes'])
                    writer.writeheader()
                    writer.writerows(medications)
                    
            elif file_name.endswith('.ics'):
                from icalendar import Calendar, Event
                from datetime import datetime, timedelta
                
                cal = Calendar()
                start_date = datetime.now()
                
                for med in medications:
                    times = []
                    if "每日一次" in med['time']:
                        times = ["08:00"]
                    elif "每日两次" in med['time']:
                        times = ["08:00", "20:00"]
                    elif "每日三次" in med['time']:
                        times = ["08:00", "14:00", "20:00"]
                    elif "每日四次" in med['time']:
                        times = ["06:00", "12:00", "18:00", "24:00"]
                    
                    for t in times:
                        event = Event()
                        event.add('summary', f"服药提醒: {med['name']}")
                        event.add('description', f"用量: {med['dosage']}\n注意事项: {med['notes']}")
                        
                        # 设置重复规则
                        event.add('rrule', {'freq': 'daily', 'count': 30})  # 30天
                        
                        # 设置提醒时间
                        hour, minute = map(int, t.split(':'))
                        event_time = start_date.replace(hour=hour, minute=minute)
                        event.add('dtstart', event_time)
                        event.add('duration', timedelta(minutes=15))
                        
                        # 添加提醒
                        event.add('alarm', timedelta(minutes=-15))
                        
                        cal.add_component(event)
                
                with open(file_name, 'wb') as f:
                    f.write(cal.to_ical())
            
            QMessageBox.information(self, "成功", "用药提醒已导出")
            
        except Exception as e:
            QMessageBox.warning(self, "警告", f"导出用药提醒失败: {str(e)}")

    def create_prescription_manager(self):
        """创建处方管理功能"""
        prescription_group = QGroupBox("处方管理")
        layout = QVBoxLayout()
        
        # 处方列表
        self.prescription_list = QTableWidget()
        self.prescription_list.setColumnCount(6)
        self.prescription_list.setHorizontalHeaderLabels([
            '处方编号', '开具日期', '患者信息', '诊断结果', '处方类型', '状态'
        ])
        
        # 设置列宽
        self.prescription_list.horizontalHeader().setSectionResizeMode(
            QHeaderView.ResizeMode.ResizeToContents)
        
        # 按钮布局
        button_layout = QHBoxLayout()
        new_btn = QPushButton("新建处方")
        edit_btn = QPushButton("编辑处方")
        delete_btn = QPushButton("删除处方")
        print_btn = QPushButton("打印处方")
        export_btn = QPushButton("导出处方")
        
        for btn in [new_btn, edit_btn, delete_btn, print_btn, export_btn]:
            button_layout.addWidget(btn)
        
        # 连接信号
        new_btn.clicked.connect(self.create_new_prescription)
        edit_btn.clicked.connect(self.edit_prescription)
        delete_btn.clicked.connect(self.delete_prescription)
        print_btn.clicked.connect(self.print_prescription)
        export_btn.clicked.connect(self.export_prescription)
        
        layout.addWidget(self.prescription_list)
        layout.addLayout(button_layout)
        prescription_group.setLayout(layout)
        
        # 初始化数据库表
        self.init_prescription_database()
        
        # 初始加载处方列表
        self.update_prescription_list()
        
        return prescription_group

    def init_prescription_database(self):
        """初始化处方数据库"""
        self.cursor.execute('''
            CREATE TABLE IF NOT EXISTS prescriptions (
                id INTEGER PRIMARY KEY AUTOINCREMENT,
                prescription_no TEXT UNIQUE,  -- 处方编号
                type TEXT,                    -- 处方类型（普通、急诊、儿科等）
                category TEXT,                -- 处方分类（西药、中药、中成药）
                date TEXT,                    -- 开具日期
                validity TEXT,                -- 有效期
                patient_name TEXT,            -- 患者姓名
                patient_gender TEXT,          -- 患者性别
                patient_age INTEGER,          -- 患者年龄
                patient_weight REAL,          -- 患者体重
                medical_insurance TEXT,       -- 医保类型
                diagnosis TEXT,               -- 诊断结果
                doctor_name TEXT,             -- 医师姓名
                doctor_title TEXT,            -- 医师职称
                hospital_name TEXT,           -- 医疗机构名称
                department TEXT,              -- 科室
                status TEXT,                  -- 状态（未调配、已调配、已发药等）
                notes TEXT                    -- 备注
            )
        ''')
        
        self.cursor.execute('''
            CREATE TABLE IF NOT EXISTS prescription_items (
                id INTEGER PRIMARY KEY AUTOINCREMENT,
                prescription_id INTEGER,       -- 关联处方ID
                medicine_name TEXT,            -- 药品名称
                specification TEXT,            -- 规格
                dosage TEXT,                   -- 用法用量
                frequency TEXT,                -- 频次
                quantity REAL,                 -- 数量
                unit TEXT,                     -- 单位
                usage_method TEXT,             -- 用药方法
                notes TEXT,                    -- 用药说明
                FOREIGN KEY (prescription_id) REFERENCES prescriptions (id)
            )
        ''')
        
        self.db.commit()

    def create_new_prescription(self):
        """创建新处方"""
        try:
            dialog = QDialog(self)
            dialog.setWindowTitle("新建处方")
            dialog.setMinimumWidth(800)
            layout = QVBoxLayout()
            
            # 处方基本信息
            basic_info = QGroupBox("基本信息")
            basic_layout = QFormLayout()
            
            # 处方类型
            prescription_type = QComboBox()
            prescription_type.addItems(['普通处方', '急诊处方', '儿科处方', '毒麻处方'])
            basic_layout.addRow("处方类型:", prescription_type)
            
            # 处方分类
            category = QComboBox()
            category.addItems(['西药处方', '中成药处方', '中药处方'])
            basic_layout.addRow("处方分类:", category)
            
            # 医疗机构信息
            hospital = QLineEdit()
            hospital.setText("XX医院")  # 默认值
            basic_layout.addRow("医疗机构:", hospital)
            
            # 科室选择
            department = QComboBox()
            department.addItems(['内科', '外科', '儿科', '妇科', '急诊科'])
            basic_layout.addRow("开具科室:", department)
            
            # 医师信息
            doctor_name = QLineEdit()
            doctor_title = QComboBox()
            doctor_title.addItems(['主任医师', '副主任医师', '主治医师', '住院医师'])
            basic_layout.addRow("医师姓名:", doctor_name)
            basic_layout.addRow("医师职称:", doctor_title)
            
            basic_info.setLayout(basic_layout)
            
            # 患者信息
            patient_info = QGroupBox("患者信息")
            patient_layout = QFormLayout()
            
            patient_name = QLineEdit()
            patient_layout.addRow("患者姓名:", patient_name)
            
            patient_gender = QComboBox()
            patient_gender.addItems(['男', '女'])
            patient_layout.addRow("性别:", patient_gender)
            
            patient_age = QSpinBox()
            patient_age.setRange(0, 120)
            patient_layout.addRow("年龄:", patient_age)
            
            patient_weight = QDoubleSpinBox()
            patient_weight.setRange(0, 200)
            patient_weight.setSuffix(" kg")
            patient_layout.addRow("体重:", patient_weight)
            
            medical_insurance = QComboBox()
            medical_insurance.addItems(['城镇职工医保', '城镇居民医保', '新农合', '自费'])
            patient_layout.addRow("医保类型:", medical_insurance)
            
            patient_info.setLayout(patient_layout)
            
            # 诊断信息
            diagnosis_info = QGroupBox("诊断信息")
            diagnosis_layout = QVBoxLayout()
            
            # 添加从分析结果提取按钮
            extract_layout = QHBoxLayout()
            extract_btn = QPushButton("从分析结果提取")
            extract_layout.addWidget(extract_btn)
            extract_layout.addStretch()
            
            diagnosis_text = QTextEdit()
            diagnosis_layout.addLayout(extract_layout)
            diagnosis_layout.addWidget(diagnosis_text)
            diagnosis_info.setLayout(diagnosis_layout)
            
            # 药品信息表格
            medicine_info = QGroupBox("药品信息")
            medicine_layout = QVBoxLayout()
            
            # 添加提取按钮
            medicine_btn_layout = QHBoxLayout()
            add_medicine_btn = QPushButton("添加药品")
            extract_medicine_btn = QPushButton("从分析结果提取")
            medicine_btn_layout.addWidget(add_medicine_btn)
            medicine_btn_layout.addWidget(extract_medicine_btn)
            medicine_btn_layout.addStretch()
            
            medicine_table = QTableWidget()
            medicine_table.setColumnCount(7)
            medicine_table.setHorizontalHeaderLabels([
                '药品名称', '规格', '用法用量', '频次', '数量', '单位', '用药说明'
            ])
            
            # 设置表格列宽
            medicine_table.horizontalHeader().setSectionResizeMode(
                QHeaderView.ResizeMode.ResizeToContents)
            
            medicine_layout.addLayout(medicine_btn_layout)
            medicine_layout.addWidget(medicine_table)
            medicine_info.setLayout(medicine_layout)
            
            # 提取诊断信息
            def extract_diagnosis():
                try:
                    # 获取分析结果文本
                    analysis_text = self.output_text.toPlainText()
                    if not analysis_text:
                        QMessageBox.warning(dialog, "警告", "没有可提取的分析结果")
                        return
                    
                    # 查找诊断部分
                    diagnosis_pattern = r"诊断[：:](.*?)(?=\n\n|$)"
                    match = re.search(diagnosis_pattern, analysis_text, re.DOTALL)
                    if match:
                        diagnosis_text.setPlainText(match.group(1).strip())
                    else:
                        QMessageBox.warning(dialog, "提示", "未找到诊断信息")
                
                except Exception as e:
                    QMessageBox.warning(dialog, "错误", f"提取诊断信息失败: {str(e)}")
            
            # 提取药品信息
            def extract_medicines():
                try:
                    # 复用用药提醒的提取功能
                    extract_dialog = QDialog(dialog)
                    extract_dialog.setWindowTitle("选择要提取的文本段落")
                    extract_layout = QVBoxLayout()
                    
                    text_edit = QTextEdit()
                    text_edit.setPlainText(self.output_text.toPlainText())
                    extract_layout.addWidget(text_edit)
                    
                    # 快速选择按钮
                    quick_select_layout = QHBoxLayout()
                    sections = ["用药方案", "治疗方案", "推荐用药", "用药建议"]
                    
                    def select_section(section_name):
                        full_text = self.output_text.toPlainText()
                        sections = full_text.split("===")
                        for section in sections:
                            if section_name in section:
                                text_edit.setPlainText(section.strip())
                                return
                    
                    for section in sections:
                        btn = QPushButton(section)
                        btn.clicked.connect(lambda checked, s=section: select_section(s))
                        quick_select_layout.addWidget(btn)
                    
                    extract_layout.addLayout(quick_select_layout)
                    
                    # 按钮
                    button_layout = QHBoxLayout()
                    extract_btn = QPushButton("提取")
                    cancel_btn = QPushButton("取消")
                    button_layout.addWidget(extract_btn)
                    button_layout.addWidget(cancel_btn)
                    extract_layout.addLayout(button_layout)
                    
                    extract_dialog.setLayout(extract_layout)
                    
                    def extract_and_add():
                        selected_text = text_edit.toPlainText()
                        medications = []
                        
                        # 使用最完整的模式进行匹配
                        pattern = r"""
                            ^-\s*([^：\n]+)：  # 药品名称
                            \s*([^，\n]+?)  # 剂量（更宽松的匹配）
                            (?:，|\s+)([^\n]+)  # 用法（允许逗号分隔）
                            (?:\s+用药说明：([^\n]+))?  # 可选的用药说明
                            (?:\s+注意事项：([^\n]+))?  # 可选的注意事项
                        """
                        
                        # 清空现有表格内容
                        medicine_table.setRowCount(0)
                        
                        matches = re.finditer(pattern, selected_text, re.MULTILINE | re.VERBOSE)
                        for match in matches:
                            if "无抗生素推荐" in match.group(1):
                                continue
                            
                            # 添加到表格
                            row = medicine_table.rowCount()
                            medicine_table.insertRow(row)
                            
                            # 药品名称
                            name_combo = QComboBox()
                            name_combo.addItems(['阿莫西林', '布洛芬', '头孢克肟', '感冒灵'])
                            name_combo.setCurrentText(match.group(1).strip())
                            medicine_table.setCellWidget(row, 0, name_combo)
                            
                            # 解析用法用量
                            dosage_parts = match.group(2).strip().split()
                            
                            # 规格
                            spec_combo = QComboBox()
                            spec_combo.addItems(['0.25g/片', '0.5g/片', '10ml/支', '5mg/片'])
                            spec_combo.setCurrentText(dosage_parts[0] if dosage_parts else '')
                            medicine_table.setCellWidget(row, 1, spec_combo)
                            
                            # 用法用量
                            dosage_combo = QComboBox()
                            dosage_combo.addItems(['1片', '2片', '5ml', '10ml'])
                            medicine_table.setCellWidget(row, 2, dosage_combo)
                            
                            # 频次
                            freq_combo = QComboBox()
                            freq_combo.addItems(['每日一次', '每日两次', '每日三次', '每4小时一次'])
                            freq_combo.setCurrentText(match.group(3).strip())
                            medicine_table.setCellWidget(row, 3, freq_combo)
                            
                            # 数量
                            quantity_spin = QSpinBox()
                            quantity_spin.setRange(1, 100)
                            quantity_spin.setValue(1)
                            medicine_table.setCellWidget(row, 4, quantity_spin)
                            
                            # 单位
                            unit_combo = QComboBox()
                            unit_combo.addItems(['片', '支', '瓶', '盒'])
                            medicine_table.setCellWidget(row, 5, unit_combo)
                            
                            # 用药说明
                            notes = ""
                            if match.group(4):  # 如果有用药说明
                                notes = f"说明：{match.group(4)}"
                                if match.group(5):  # 如果有注意事项
                                    notes += f"\n注意：{match.group(5)}"
                            notes_edit = QLineEdit()
                            notes_edit.setText(notes)
                            medicine_table.setCellWidget(row, 6, notes_edit)
                        
                        # 如果没有匹配到任何药品，尝试使用简单模式
                        if medicine_table.rowCount() == 0:
                            pattern = r"-\s*([^：\n]+)：\s*([^，\n]+?)(?:，|\s+)([^\n]+)"
                            matches = re.finditer(pattern, selected_text, re.MULTILINE)
                            for match in matches:
                                if "无抗生素推荐" in match.group(1):
                                    continue
                                
                                # 添加到表格（使用相同的添加逻辑）
                                row = medicine_table.rowCount()
                                medicine_table.insertRow(row)
                                # ... (添加药品信息的代码与上面相同)
                        
                        extract_dialog.accept()
                    
                    extract_btn.clicked.connect(extract_and_add)
                    cancel_btn.clicked.connect(extract_dialog.reject)
                    
                    extract_dialog.exec()
                    
                except Exception as e:
                    QMessageBox.warning(dialog, "错误", f"提取药品信息失败: {str(e)}")
            
            # 连接信号
            extract_btn.clicked.connect(extract_diagnosis)
            extract_medicine_btn.clicked.connect(extract_medicines)
            
            # 添加所有组件到主布局
            for widget in [basic_info, patient_info, diagnosis_info, medicine_info]:
                layout.addWidget(widget)
            
            # 底部按钮
            button_layout = QHBoxLayout()
            save_btn = QPushButton("保存处方")
            cancel_btn = QPushButton("取消")
            
            def save_prescription():
                try:
                    # 生成处方编号
                    prescription_no = f"RX{datetime.now().strftime('%Y%m%d%H%M%S')}"
                    
                    # 保存处方基本信息
                    self.cursor.execute('''
                        INSERT INTO prescriptions (
                            prescription_no, type, category, date, patient_name,
                            patient_gender, patient_age, patient_weight,
                            medical_insurance, diagnosis, doctor_name,
                            doctor_title, hospital_name, department, status
                        ) VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?)
                    ''', (
                        prescription_no,
                        prescription_type.currentText(),
                        category.currentText(),
                        datetime.now().strftime('%Y-%m-%d %H:%M:%S'),
                        patient_name.text(),
                        patient_gender.currentText(),
                        patient_age.value(),
                        patient_weight.value(),
                        medical_insurance.currentText(),
                        diagnosis_text.toPlainText(),
                        doctor_name.text(),
                        doctor_title.currentText(),
                        hospital.text(),
                        department.currentText(),
                        "未调配"
                    ))
                    
                    prescription_id = self.cursor.lastrowid
                    
                    # 保存药品信息
                    for row in range(medicine_table.rowCount()):
                        name_widget = medicine_table.cellWidget(row, 0)
                        spec_widget = medicine_table.cellWidget(row, 1)
                        dosage_widget = medicine_table.cellWidget(row, 2)
                        freq_widget = medicine_table.cellWidget(row, 3)
                        quantity_widget = medicine_table.cellWidget(row, 4)
                        unit_widget = medicine_table.cellWidget(row, 5)
                        notes_widget = medicine_table.cellWidget(row, 6)
                        
                        if name_widget:
                            self.cursor.execute('''
                                INSERT INTO prescription_items (
                                    prescription_id, medicine_name, specification,
                                    dosage, frequency, quantity, unit, usage_method
                                ) VALUES (?, ?, ?, ?, ?, ?, ?, ?)
                            ''', (
                                prescription_id,
                                name_widget.currentText(),
                                spec_widget.currentText(),
                                dosage_widget.currentText(),
                                freq_widget.currentText(),
                                quantity_widget.value(),
                                unit_widget.currentText(),
                                notes_widget.text()
                            ))
                    
                    self.db.commit()
                    
                    # 更新处方列表
                    self.update_prescription_list()
                    
                    QMessageBox.information(dialog, "成功", "处方保存成功")
                    dialog.accept()
                    
                except Exception as e:
                    QMessageBox.warning(dialog, "错误", f"保存处方失败: {str(e)}")
            
            save_btn.clicked.connect(save_prescription)
            cancel_btn.clicked.connect(dialog.reject)
            
            button_layout.addWidget(save_btn)
            button_layout.addWidget(cancel_btn)
            layout.addLayout(button_layout)
            
            dialog.setLayout(layout)
            dialog.exec()
            
        except Exception as e:
            QMessageBox.warning(self, "错误", f"创建处方失败: {str(e)}")

    def update_prescription_list(self):
        """更新处方列表显示"""
        try:
            # 清空现有列表
            self.prescription_list.setRowCount(0)
            
            # 获取所有处方信息
            self.cursor.execute('''
                SELECT 
                    prescription_no, date, 
                    patient_name || ' ' || patient_age || '岁 ' || patient_gender as patient_info,
                    diagnosis, type, status
                FROM prescriptions 
                ORDER BY date DESC
            ''')
            
            prescriptions = self.cursor.fetchall()
            
            # 添加到表格
            for row, prescription in enumerate(prescriptions):
                self.prescription_list.insertRow(row)
                for col, value in enumerate(prescription):
                    self.prescription_list.setItem(row, col, QTableWidgetItem(str(value)))
            
        except Exception as e:
            QMessageBox.warning(self, "错误", f"更新处方列表失败: {str(e)}")

    def edit_prescription(self):
        """编辑处方"""
        current_row = self.prescription_list.currentRow()
        if current_row < 0:
            QMessageBox.warning(self, "警告", "请先选择要编辑的处方")
            return
        
        try:
            # 获取处方信息
            prescription_no = self.prescription_list.item(current_row, 0).text()
            
            # 从数据库获取详细信息
            self.cursor.execute('''
                SELECT * FROM prescriptions WHERE prescription_no = ?
            ''', (prescription_no,))
            prescription = self.cursor.fetchone()
            
            if not prescription:
                QMessageBox.warning(self, "错误", "未找到处方信息")
                return
            
            # 获取药品信息
            self.cursor.execute('''
                SELECT * FROM prescription_items WHERE prescription_id = ?
            ''', (prescription[0],))
            items = self.cursor.fetchall()
            
            # 创建编辑对话框
            dialog = QDialog(self)
            dialog.setWindowTitle("编辑处方")
            dialog.setMinimumWidth(800)
            layout = QVBoxLayout()
            
            # 基本信息
            basic_info = QGroupBox("基本信息")
            basic_layout = QFormLayout()
            
            # 处方类型
            prescription_type = QComboBox()
            prescription_type.addItems(['普通处方', '急诊处方', '儿科处方', '毒麻处方'])
            prescription_type.setCurrentText(prescription[2])  # type
            basic_layout.addRow("处方类型:", prescription_type)
            
            # 处方分类
            category = QComboBox()
            category.addItems(['西药处方', '中成药处方', '中药处方'])
            category.setCurrentText(prescription[3])  # category
            basic_layout.addRow("处方分类:", category)
            
            # 医疗机构
            hospital = QLineEdit()
            hospital.setText(prescription[14])  # hospital_name
            basic_layout.addRow("医疗机构:", hospital)
            
            # 科室
            department = QComboBox()
            department.addItems(['内科', '外科', '儿科', '妇科', '急诊科'])
            department.setCurrentText(prescription[15])  # department
            basic_layout.addRow("开具科室:", department)
            
            # 医师信息
            doctor_name = QLineEdit()
            doctor_name.setText(prescription[12])  # doctor_name
            doctor_title = QComboBox()
            doctor_title.addItems(['主任医师', '副主任医师', '主治医师', '住院医师'])
            doctor_title.setCurrentText(prescription[13])  # doctor_title
            basic_layout.addRow("医师姓名:", doctor_name)
            basic_layout.addRow("医师职称:", doctor_title)
            
            basic_info.setLayout(basic_layout)
            
            # 患者信息
            patient_info = QGroupBox("患者信息")
            patient_layout = QFormLayout()
            
            patient_name = QLineEdit()
            patient_name.setText(prescription[6])  # patient_name
            patient_layout.addRow("患者姓名:", patient_name)
            
            patient_gender = QComboBox()
            patient_gender.addItems(['男', '女'])
            patient_gender.setCurrentText(prescription[7])  # patient_gender
            patient_layout.addRow("性别:", patient_gender)
            
            patient_age = QSpinBox()
            patient_age.setRange(0, 120)
            patient_age.setValue(prescription[8])  # patient_age
            patient_layout.addRow("年龄:", patient_age)
            
            patient_weight = QDoubleSpinBox()
            patient_weight.setRange(0, 200)
            patient_weight.setSuffix(" kg")
            patient_weight.setValue(prescription[9])  # patient_weight
            patient_layout.addRow("体重:", patient_weight)
            
            medical_insurance = QComboBox()
            medical_insurance.addItems(['城镇职工医保', '城镇居民医保', '新农合', '自费'])
            medical_insurance.setCurrentText(prescription[10])  # medical_insurance
            patient_layout.addRow("医保类型:", medical_insurance)
            
            patient_info.setLayout(patient_layout)
            
            # 诊断信息
            diagnosis_info = QGroupBox("诊断信息")
            diagnosis_layout = QVBoxLayout()
            diagnosis_text = QTextEdit()
            diagnosis_text.setPlainText(prescription[11])  # diagnosis
            diagnosis_layout.addWidget(diagnosis_text)
            diagnosis_info.setLayout(diagnosis_layout)
            
            # 药品信息
            medicine_info = QGroupBox("药品信息")
            medicine_layout = QVBoxLayout()
            
            medicine_table = QTableWidget()
            medicine_table.setColumnCount(7)
            medicine_table.setHorizontalHeaderLabels([
                '药品名称', '规格', '用法用量', '频次', '数量', '单位', '用药说明'
            ])
            
            # 添加现有药品
            for item in items:
                row = medicine_table.rowCount()
                medicine_table.insertRow(row)
                
                # 药品名称
                name_combo = QComboBox()
                name_combo.addItems(['阿莫西林', '布洛芬', '头孢克肟', '感冒灵'])
                name_combo.setCurrentText(item[2])  # medicine_name
                medicine_table.setCellWidget(row, 0, name_combo)
                
                # 规格
                spec_combo = QComboBox()
                spec_combo.addItems(['0.25g/片', '0.5g/片', '10ml/支', '5mg/片'])
                spec_combo.setCurrentText(item[3])  # specification
                medicine_table.setCellWidget(row, 1, spec_combo)
                
                # 用法用量
                dosage_combo = QComboBox()
                dosage_combo.addItems(['1片', '2片', '5ml', '10ml'])
                dosage_combo.setCurrentText(item[4])  # dosage
                medicine_table.setCellWidget(row, 2, dosage_combo)
                
                # 频次
                freq_combo = QComboBox()
                freq_combo.addItems(['每日一次', '每日两次', '每日三次', '每4小时一次'])
                freq_combo.setCurrentText(item[5])  # frequency
                medicine_table.setCellWidget(row, 3, freq_combo)
                
                # 数量
                quantity_spin = QSpinBox()
                quantity_spin.setRange(1, 100)
                quantity_spin.setValue(int(item[6]))  # quantity
                medicine_table.setCellWidget(row, 4, quantity_spin)
                
                # 单位
                unit_combo = QComboBox()
                unit_combo.addItems(['片', '支', '瓶', '盒'])
                unit_combo.setCurrentText(item[7])  # unit
                medicine_table.setCellWidget(row, 5, unit_combo)
                
                # 用药说明
                notes_edit = QLineEdit()
                notes_edit.setText(item[9])  # notes
                medicine_table.setCellWidget(row, 6, notes_edit)
            
            # 添加药品按钮
            add_btn = QPushButton("添加药品")
            def add_medicine_row():
                row = medicine_table.rowCount()
                medicine_table.insertRow(row)
                
                # 添加新的控件
                name_combo = QComboBox()
                name_combo.addItems(['阿莫西林', '布洛芬', '头孢克肟', '感冒灵'])
                medicine_table.setCellWidget(row, 0, name_combo)
                
                spec_combo = QComboBox()
                spec_combo.addItems(['0.25g/片', '0.5g/片', '10ml/支', '5mg/片'])
                medicine_table.setCellWidget(row, 1, spec_combo)
                
                dosage_combo = QComboBox()
                dosage_combo.addItems(['1片', '2片', '5ml', '10ml'])
                medicine_table.setCellWidget(row, 2, dosage_combo)
                
                freq_combo = QComboBox()
                freq_combo.addItems(['每日一次', '每日两次', '每日三次', '每4小时一次'])
                medicine_table.setCellWidget(row, 3, freq_combo)
                
                quantity_spin = QSpinBox()
                quantity_spin.setRange(1, 100)
                medicine_table.setCellWidget(row, 4, quantity_spin)
                
                unit_combo = QComboBox()
                unit_combo.addItems(['片', '支', '瓶', '盒'])
                medicine_table.setCellWidget(row, 5, unit_combo)
                
                notes_edit = QLineEdit()
                medicine_table.setCellWidget(row, 6, notes_edit)
            
            add_btn.clicked.connect(add_medicine_row)
            
            medicine_layout.addWidget(medicine_table)
            medicine_layout.addWidget(add_btn)
            medicine_info.setLayout(medicine_layout)
            
            # 添加所有组件
            for widget in [basic_info, patient_info, diagnosis_info, medicine_info]:
                layout.addWidget(widget)
            
            # 底部按钮
            button_layout = QHBoxLayout()
            save_btn = QPushButton("保存")
            cancel_btn = QPushButton("取消")
            
            def save_changes():
                try:
                    # 更新处方基本信息
                    self.cursor.execute('''
                        UPDATE prescriptions SET
                        type = ?, category = ?, hospital_name = ?, department = ?,
                        doctor_name = ?, doctor_title = ?, patient_name = ?,
                        patient_gender = ?, patient_age = ?, patient_weight = ?,
                        medical_insurance = ?, diagnosis = ?
                        WHERE prescription_no = ?
                    ''', (
                        prescription_type.currentText(),
                        category.currentText(),
                        hospital.text(),
                        department.currentText(),
                        doctor_name.text(),
                        doctor_title.currentText(),
                        patient_name.text(),
                        patient_gender.currentText(),
                        patient_age.value(),
                        patient_weight.value(),
                        medical_insurance.currentText(),
                        diagnosis_text.toPlainText(),
                        prescription_no
                    ))
                    
                    # 删除原有药品信息
                    self.cursor.execute('''
                        DELETE FROM prescription_items WHERE prescription_id = ?
                    ''', (prescription[0],))
                    
                    # 添加新的药品信息
                    for row in range(medicine_table.rowCount()):
                        name_widget = medicine_table.cellWidget(row, 0)
                        spec_widget = medicine_table.cellWidget(row, 1)
                        dosage_widget = medicine_table.cellWidget(row, 2)
                        freq_widget = medicine_table.cellWidget(row, 3)
                        quantity_widget = medicine_table.cellWidget(row, 4)
                        unit_widget = medicine_table.cellWidget(row, 5)
                        notes_widget = medicine_table.cellWidget(row, 6)
                        
                        if name_widget:
                            self.cursor.execute('''
                                INSERT INTO prescription_items (
                                    prescription_id, medicine_name, specification,
                                    dosage, frequency, quantity, unit, usage_method
                                ) VALUES (?, ?, ?, ?, ?, ?, ?, ?)
                            ''', (
                                prescription[0],
                                name_widget.currentText(),
                                spec_widget.currentText(),
                                dosage_widget.currentText(),
                                freq_widget.currentText(),
                                quantity_widget.value(),
                                unit_widget.currentText(),
                                notes_widget.text()
                            ))
                    
                    self.db.commit()
                    
                    # 更新处方列表
                    self.update_prescription_list()
                    
                    QMessageBox.information(dialog, "成功", "处方已更新")
                    dialog.accept()
                    
                except Exception as e:
                    QMessageBox.warning(dialog, "错误", f"保存更改失败: {str(e)}")
            
            save_btn.clicked.connect(save_changes)
            cancel_btn.clicked.connect(dialog.reject)
            
            button_layout.addWidget(save_btn)
            button_layout.addWidget(cancel_btn)
            layout.addLayout(button_layout)
            
            dialog.setLayout(layout)
            dialog.exec()
            
        except Exception as e:
            QMessageBox.warning(self, "错误", f"编辑处方失败: {str(e)}")

    def delete_prescription(self):
        """删除处方"""
        current_row = self.prescription_list.currentRow()
        if current_row < 0:
            QMessageBox.warning(self, "警告", "请先选择要删除的处方")
            return
        
        try:
            prescription_no = self.prescription_list.item(current_row, 0).text()
            
            if QMessageBox.question(
                self,
                "确认删除",
                f"确定要删除处方 {prescription_no} 吗？",
                QMessageBox.StandardButton.Yes | QMessageBox.StandardButton.No
            ) == QMessageBox.StandardButton.Yes:  # 修改这里
                # 删除处方及相关药品信息
                self.cursor.execute('''
                    DELETE FROM prescription_items 
                    WHERE prescription_id IN (
                        SELECT id FROM prescriptions WHERE prescription_no = ?
                    )
                ''', (prescription_no,))
                
                self.cursor.execute('''
                    DELETE FROM prescriptions WHERE prescription_no = ?
                ''', (prescription_no,))
                
                self.db.commit()
                
                # 从表格中删除
                self.prescription_list.removeRow(current_row)
                
                QMessageBox.information(self, "成功", "处方已删除")
        
        except Exception as e:
            QMessageBox.warning(self, "错误", f"删除处方失败: {str(e)}")

    def print_prescription(self):
        """打印处方"""
        current_row = self.prescription_list.currentRow()
        if current_row < 0:
            QMessageBox.warning(self, "警告", "请先选择要打印的处方")
            return
        
        try:
            prescription_no = self.prescription_list.item(current_row, 0).text()
            
            # 获取处方信息
            self.cursor.execute('''
                SELECT p.*, i.* 
                FROM prescriptions p
                LEFT JOIN prescription_items i ON p.id = i.prescription_id
                WHERE p.prescription_no = ?
            ''', (prescription_no,))
            
            prescription_data = self.cursor.fetchall()
            
            if not prescription_data:
                QMessageBox.warning(self, "错误", "未找到处方信息")
                return
            
            # 创建打印预览
            dialog = QPrintPreviewDialog()
            dialog.paintRequested.connect(lambda printer: self.print_prescription_content(printer, prescription_data))
            dialog.exec()
            
        except Exception as e:
            QMessageBox.warning(self, "错误", f"打印处方失败: {str(e)}")

    def print_prescription_content(self, printer, prescription_data):
        """打印处方内容"""
        try:
            painter = QPainter()
            painter.begin(printer)
            
            # 设置字体
            font = QFont("SimSun", 12)  # 使用宋体
            painter.setFont(font)
            
            # 打印处方内容
            y = 100
            
            # 标题
            font.setPointSize(16)
            painter.setFont(font)
            painter.drawText(300, y, "处方笺")
            
            # 恢复字体大小
            font.setPointSize(12)
            painter.setFont(font)
            
            y += 50
            
            # 打印基本信息
            basic_info = prescription_data[0]
            painter.drawText(100, y, f"处方编号: {basic_info[1]}")
            y += 30
            painter.drawText(100, y, f"开具日期: {basic_info[4]}")
            y += 30
            painter.drawText(100, y, f"医疗机构: {basic_info[14]}")
            y += 30
            painter.drawText(100, y, f"科室: {basic_info[15]}")
            y += 30
            
            # 患者信息
            y += 20
            painter.drawText(100, y, "患者信息:")
            y += 30
            painter.drawText(120, y, f"姓名: {basic_info[6]}    性别: {basic_info[7]}    年龄: {basic_info[8]}岁")
            y += 30
            painter.drawText(120, y, f"体重: {basic_info[9]}kg    医保类型: {basic_info[10]}")
            
            # 诊断信息
            y += 50
            painter.drawText(100, y, "诊断:")
            y += 30
            
            # 处理多行诊断
            diagnosis_lines = basic_info[11].split('\n')
            for line in diagnosis_lines:
                painter.drawText(120, y, line)
                y += 20
            
            # 药品信息
            y += 30
            painter.drawText(100, y, "处方药品:")
            y += 30
            
            # 表头
            painter.drawText(120, y, "药品名称")
            painter.drawText(300, y, "规格")
            painter.drawText(400, y, "用法用量")
            painter.drawText(500, y, "数量")
            y += 20
            
            # 药品列表
            for item in prescription_data:
                if item[19]:  # 如果有药品信息
                    painter.drawText(120, y, item[20])  # 药品名称
                    painter.drawText(300, y, item[21])  # 规格
                    painter.drawText(400, y, f"{item[22]} {item[23]}")  # 用法用量和频次
                    painter.drawText(500, y, f"{item[24]}{item[25]}")  # 数量和单位
                    y += 30
                    if item[27]:  # 如果有用药说明
                        painter.drawText(140, y, f"说明: {item[27]}")
                        y += 20
            
            # 医师信息
            y += 50
            painter.drawText(100, y, f"医师: {basic_info[12]}    职称: {basic_info[13]}")
            
            # 签名和日期
            y += 50
            painter.drawText(100, y, "医师签名: _____________    日期: _____________")
            
            # 注意事项
            y += 50
            painter.drawText(100, y, "注意事项:")
            y += 20
            painter.drawText(120, y, "1. 请按医嘱用药，不得擅自加减药量或停药")
            y += 20
            painter.drawText(120, y, "2. 如有不适，请及时就医")
            y += 20
            painter.drawText(120, y, "3. 本处方仅限本次使用")
            
            painter.end()
            
        except Exception as e:
            QMessageBox.warning(self, "错误", f"生成打印内容失败: {str(e)}")

    def export_prescription(self):
        """导出处方"""
        current_row = self.prescription_list.currentRow()
        if current_row < 0:
            QMessageBox.warning(self, "警告", "请先选择要导出的处方")
            return
        
        try:
            prescription_no = self.prescription_list.item(current_row, 0).text()
            
            file_name, _ = QFileDialog.getSaveFileName(
                self,
                "导出处方",
                f"处方_{prescription_no}",
                "PDF文件 (*.pdf);;Word文档 (*.docx)"
            )
            
            if not file_name:
                return
            
            # 获取处方信息
            self.cursor.execute('''
                SELECT p.*, i.* 
                FROM prescriptions p
                LEFT JOIN prescription_items i ON p.id = i.prescription_id
                WHERE p.prescription_no = ?
            ''', (prescription_no,))
            
            prescription_data = self.cursor.fetchall()
            
            if file_name.endswith('.pdf'):
                self.export_prescription_pdf(file_name, prescription_data)
            else:
                self.export_prescription_docx(file_name, prescription_data)
            
            QMessageBox.information(self, "成功", "处方已导出")
            
        except Exception as e:
            QMessageBox.warning(self, "错误", f"导出处方失败: {str(e)}")

    def export_prescription_pdf(self, file_name: str, prescription_data: list):
        """导出处方为PDF"""
        from reportlab.pdfgen import canvas
        from reportlab.lib.pagesizes import A4
        from reportlab.pdfbase import pdfmetrics
        from reportlab.pdfbase.ttfonts import TTFont
        
        # 根据操作系统选择合适的中文字体
        system = platform.system()
        if system == "Darwin":  # macOS
            try:
                # 尝试使用苹果系统自带的中文字体
                font_path = "/System/Library/Fonts/PingFang.ttc"
                font_name = "PingFang"
                pdfmetrics.registerFont(TTFont(font_name, font_path))
            except:
                # 如果找不到 PingFang，使用其他可能存在的字体
                try:
                    font_path = "/System/Library/Fonts/STHeiti Light.ttc"
                    font_name = "STHeiti"
                    pdfmetrics.registerFont(TTFont(font_name, font_path))
                except:
                    # 如果都找不到，使用默认字体
                    font_name = "Helvetica"
        elif system == "Windows":
            try:
                font_path = "C:\\Windows\\Fonts\\simsun.ttc"
                font_name = "SimSun"
                pdfmetrics.registerFont(TTFont(font_name, font_path))
            except:
                font_name = "Helvetica"
        else:
            font_name = "Helvetica"
        
        c = canvas.Canvas(file_name, pagesize=A4)
        c.setFont(font_name, 12)
        
        # 添加处方内容
        y = 800
        c.setFont(font_name, 16)
        c.drawString(250, y, "处方笺")
        y -= 50
        
        c.setFont(font_name, 12)
        # 基本信息
        basic_info = prescription_data[0]
        c.drawString(100, y, f"处方编号: {basic_info[1]}")
        y -= 30
        c.drawString(100, y, f"开具日期: {basic_info[4]}")
        y -= 30
        c.drawString(100, y, f"医疗机构: {basic_info[14]}")
        y -= 30
        c.drawString(100, y, f"科室: {basic_info[15]}")
        y -= 30
        
        # 患者信息
        y -= 20
        c.drawString(100, y, "患者信息:")
        y -= 30
        c.drawString(120, y, f"姓名: {basic_info[6]}    性别: {basic_info[7]}    年龄: {basic_info[8]}岁")
        y -= 30
        c.drawString(120, y, f"体重: {basic_info[9]}kg    医保类型: {basic_info[10]}")
        
        # 诊断信息
        y -= 40
        c.drawString(100, y, "诊断:")
        y -= 30
        
        # 处理多行诊断
        diagnosis_lines = basic_info[11].split('\n')
        for line in diagnosis_lines:
            c.drawString(120, y, line)
            y -= 20
        
        # 药品信息
        y -= 30
        c.drawString(100, y, "处方药品:")
        y -= 30
        
        # 表头
        c.drawString(120, y, "药品名称")
        c.drawString(300, y, "规格")
        c.drawString(400, y, "用法用量")
        c.drawString(500, y, "数量")
        y -= 20
        
        # 药品列表
        for item in prescription_data:
            if item[19]:  # 如果有药品信息
                c.drawString(120, y, item[20])  # 药品名称
                c.drawString(300, y, item[21])  # 规格
                c.drawString(400, y, f"{item[22]} {item[23]}")  # 用法用量和频次
                c.drawString(500, y, f"{item[24]}{item[25]}")  # 数量和单位
                y -= 30
                if item[27]:  # 如果有用药说明
                    c.drawString(140, y, f"说明: {item[27]}")
                    y -= 20
        
        # 医师信息
        y -= 40
        c.drawString(100, y, f"医师: {basic_info[12]}    职称: {basic_info[13]}")
        
        # 签名和日期
        y -= 40
        c.drawString(100, y, "医师签名: _____________    日期: _____________")
        
        # 注意事项
        y -= 40
        c.drawString(100, y, "注意事项:")
        y -= 20
        c.drawString(120, y, "1. 请按医嘱用药，不得擅自加减药量或停药")
        y -= 20
        c.drawString(120, y, "2. 如有不适，请及时就医")
        y -= 20
        c.drawString(120, y, "3. 本处方仅限本次使用")
        
        c.save()

    def export_prescription_docx(self, file_name: str, prescription_data: list):
        """导出处方为Word文档"""
        from docx import Document
        
        doc = Document()
        doc.add_heading('处方', 0)
        
        # 添加基本信息
        basic_info = prescription_data[0]
        doc.add_paragraph(f"处方编号: {basic_info[1]}")
        doc.add_paragraph(f"开具日期: {basic_info[4]}")
        doc.add_paragraph(f"患者姓名: {basic_info[6]}")
        
        # 添加药品信息
        doc.add_heading('药品信息', level=1)
        for item in prescription_data:
            if item[19]:  # 如果有药品信息
                doc.add_paragraph(f"{item[20]} {item[21]} {item[22]}")
        
        doc.save(file_name)

    def update_health_chart(self):
        """更新健康趋势图表"""
        try:
            import matplotlib.pyplot as plt
            import matplotlib.dates as mdates
            trend_type = self.trend_type.currentText()
            time_range = self.range_combo.currentText()
            
            # 根据时间范围构建查询
            if time_range == '最近7天':
                days = 7
            elif time_range == '最近30天':
                days = 30
            elif time_range == '最近90天':
                days = 90
            else:
                days = None
            
            if days:
                query = '''
                    SELECT timestamp, value 
                    FROM health_trends 
                    WHERE type = ? AND timestamp >= datetime('now', ?)
                    ORDER BY timestamp
                '''
                params = (trend_type, f'-{days} days')
            else:
                query = '''
                    SELECT timestamp, value 
                    FROM health_trends 
                    WHERE type = ?
                    ORDER BY timestamp
                '''
                params = (trend_type,)
            
            # 获取数据
            self.cursor.execute(query, params)
            data = self.cursor.fetchall()
            
            # 清除当前图表
            self.ax.clear()
            
            if data:
                # 转换数据
                dates = [datetime.strptime(row[0], '%Y-%m-%d %H:%M:%S') for row in data]
                values = [row[1] for row in data]
                
                # 绘制图表
                self.ax.plot(dates, values, marker='o', linestyle='-', color='b', label=trend_type)
                
                # 设置标签
                self.ax.set_title(f'{trend_type}趋势图')
                self.ax.set_xlabel('日期')
                self.ax.set_ylabel(trend_type)
                self.ax.legend()  # 添加图例
                
                # 设置x轴日期格式
                self.ax.xaxis.set_major_formatter(mdates.DateFormatter('%Y-%m-%d'))
                self.figure.autofmt_xdate()  # 自动旋转日期标签
                
                # 添加网格
                self.ax.grid(True, linestyle='--', alpha=0.7)
                
                # 设置合适的刻度
                self.ax.yaxis.set_major_locator(plt.MaxNLocator(10))
            else:
                self.ax.text(0.5, 0.5, '暂无数据', 
                            horizontalalignment='center',
                            verticalalignment='center',
                            transform=self.ax.transAxes)
            
            # 刷新画布
            self.canvas.draw()
            
        except Exception as e:
            print(f"更新图表失败: {str(e)}")

    def show_tutorial(self):
        QMessageBox.information(self, "欢迎使用", "这是一个AI 医疗助手应用- AI安全工坊出品（微信公众号搜索关注），您可以通过以下功能进行操作：\n1. 输入患者信息\n2. 描述症状\n3. 获取诊断结果\n4. 管理用药提醒\n5. 查看健康趋势")

    def create_search_bar(self):
        search_bar = QLineEdit()
        search_bar.setPlaceholderText("搜索病历或处方...")
        search_bar.textChanged.connect(self.filter_records)
        return search_bar

    def filter_records(self, text):
        """根据输入的 text 过滤病历记录"""
        for row in range(self.patient_record_table.rowCount()):
            item = self.patient_record_table.item(row, 1)  # 假设患者信息在第二列
            if item and text.lower() in item.text().lower():
                self.patient_record_table.showRow(row)
            else:
                self.patient_record_table.hideRow(row)

    def create_patient_record_table(self):
        """创建病历记录表格"""
        self.patient_record_table = QTableWidget()
        self.patient_record_table.setColumnCount(4)
        self.patient_record_table.setHorizontalHeaderLabels(['时间', '患者信息', '症状', '诊断'])
        
        # 添加搜索框
        self.search_bar = self.create_search_bar()
        
        # 布局
        layout = QVBoxLayout()
        layout.addWidget(self.search_bar)
        layout.addWidget(self.patient_record_table)
        
        # 设置布局
        group = QGroupBox("病历记录")
        group.setLayout(layout)
        return group

if __name__ == '__main__':
    app = QApplication(sys.argv)
    
    # 设置应用样式
    app.setStyle("Fusion")
    
    # 创建并显示主窗口
    window = MedicalAssistant()
    window.show()
    
    sys.exit(app.exec()) 